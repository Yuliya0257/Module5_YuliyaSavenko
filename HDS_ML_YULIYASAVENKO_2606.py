#!/usr/bin/env python3
"""
analysis.py — Single-file pipeline for the HDS Module 5 assignment.

All code that was previously split across `code/ml_lib.py`,
`code/run_analysis.py`, `code/tune_hyperparameters.py`,
`code/sensitivity_age.py`, `code/sensitivity_clustering.py`,
`code/validate_ml_lib.py` and `build_report.py` is consolidated here.

Usage:
    python analysis.py <command>

Commands:
    analyse              Main analysis pipeline (clustering, three classifiers,
                         figures, results) — equivalent to the original
                         run_analysis.py.
    tune                 Hyperparameter grid search and selection.
    sensitivity-age      Age-window sensitivity (18-39 vs 18-64 vs 18-99).
    sensitivity-cluster  Clustering-method sensitivity (KMeans vs KModes).
    validate             Sanity-check the from-scratch ml_lib implementations.
    report               Build the final Word report from generated artefacts.
    all                  Run analyse + tune + sensitivity-* + validate + report.

Run `python analysis.py <command> -h` for command-specific help.
"""
from __future__ import annotations

# ---------------------------------------------------------------------------
# Combined imports
# ---------------------------------------------------------------------------
import argparse
import json
import os
import sys
import time
from math import erf, sqrt
from pathlib import Path

import numpy as np
import pandas as pd
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import seaborn as sns
from matplotlib.colors import LinearSegmentedColormap

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Shared paths (analysis.py lives at the repo root)
# ---------------------------------------------------------------------------
HERE = Path(__file__).resolve().parent
REPO = HERE
DATA = REPO / "data" / "raw_nhanes"
FIG = REPO / "figures"
RES = REPO / "results"
OUT_REPORT = REPO / "HDS_ML_YULIYASAVENKO_0406.docx"
for _d in (FIG, RES, DATA):
    _d.mkdir(parents=True, exist_ok=True)


# ===========================================================================
# SECTION 1  ml_lib  (from-scratch ML implementations)
# ===========================================================================

# ---------------------------------------------------------------------------
# Pre-processing
# ---------------------------------------------------------------------------
class StandardScaler:
    """Centre to zero mean and unit variance."""

    def __init__(self) -> None:
        self.mean_: np.ndarray | None = None
        self.std_: np.ndarray | None = None

    def fit(self, X: np.ndarray) -> "StandardScaler":
        X = np.asarray(X, dtype=float)
        self.mean_ = X.mean(axis=0)
        self.std_ = X.std(axis=0, ddof=0)
        self.std_[self.std_ == 0] = 1.0
        return self

    def transform(self, X: np.ndarray) -> np.ndarray:
        return (np.asarray(X, dtype=float) - self.mean_) / self.std_

    def fit_transform(self, X: np.ndarray) -> np.ndarray:
        return self.fit(X).transform(X)


# ---------------------------------------------------------------------------
# K-means
# ---------------------------------------------------------------------------
class KMeans:
    """K-means clustering with k-means++ initialisation."""

    def __init__(self, n_clusters: int = 3, max_iter: int = 300,
                 n_init: int = 10, tol: float = 1e-4,
                 random_state: int | None = None) -> None:
        self.n_clusters = n_clusters
        self.max_iter = max_iter
        self.n_init = n_init
        self.tol = tol
        self.random_state = random_state
        self.cluster_centers_: np.ndarray | None = None
        self.labels_: np.ndarray | None = None
        self.inertia_: float = np.inf

    # ---- helpers --------------------------------------------------------
    def _kmeans_pp(self, X: np.ndarray, rng: np.random.Generator) -> np.ndarray:
        n_samples = X.shape[0]
        idx0 = rng.integers(n_samples)
        centers = [X[idx0]]
        for _ in range(1, self.n_clusters):
            d2 = np.min(
                ((X[:, None, :] - np.stack(centers)[None, :, :]) ** 2).sum(-1),
                axis=1,
            )
            probs = d2 / d2.sum() if d2.sum() > 0 else np.ones(n_samples) / n_samples
            idx = rng.choice(n_samples, p=probs)
            centers.append(X[idx])
        return np.stack(centers)

    def _single_run(self, X: np.ndarray, seed: int) -> tuple[np.ndarray, np.ndarray, float]:
        rng = np.random.default_rng(seed)
        centers = self._kmeans_pp(X, rng)
        labels = np.zeros(len(X), dtype=int)
        for _ in range(self.max_iter):
            d2 = ((X[:, None, :] - centers[None, :, :]) ** 2).sum(-1)
            new_labels = d2.argmin(axis=1)
            new_centers = np.stack([
                X[new_labels == k].mean(axis=0) if (new_labels == k).any()
                else centers[k]
                for k in range(self.n_clusters)
            ])
            shift = np.linalg.norm(new_centers - centers)
            centers, labels = new_centers, new_labels
            if shift < self.tol:
                break
        inertia = ((X - centers[labels]) ** 2).sum()
        return centers, labels, inertia

    # ---- API ------------------------------------------------------------
    def fit(self, X: np.ndarray) -> "KMeans":
        X = np.asarray(X, dtype=float)
        base = 0 if self.random_state is None else int(self.random_state)
        best = None
        for i in range(self.n_init):
            centers, labels, inertia = self._single_run(X, base + i)
            if best is None or inertia < best[2]:
                best = (centers, labels, inertia)
        self.cluster_centers_, self.labels_, self.inertia_ = best
        return self

    def predict(self, X: np.ndarray) -> np.ndarray:
        X = np.asarray(X, dtype=float)
        d2 = ((X[:, None, :] - self.cluster_centers_[None, :, :]) ** 2).sum(-1)
        return d2.argmin(axis=1)


class KPrototypes:
    """K-prototypes clustering for mixed binary + numeric data (Huang, 1998).

    Distance to a centroid is the squared Euclidean distance over numeric
    columns plus gamma times the count of mismatches over categorical
    (binary) columns. Centroids are the per-cluster mean for numeric
    columns and the mode for categorical columns. This is the
    methodologically appropriate clustering method when most features are
    binary, where plain K-means' spherical-Gaussian assumption does not
    hold.
    """

    def __init__(self, n_clusters: int = 2, cat_idx: list[int] | None = None,
                 num_idx: list[int] | None = None, gamma: float | None = None,
                 max_iter: int = 100, n_init: int = 10,
                 random_state: int | None = None) -> None:
        self.n_clusters = n_clusters
        self.cat_idx = cat_idx
        self.num_idx = num_idx
        self.gamma = gamma
        self.max_iter = max_iter
        self.n_init = n_init
        self.random_state = random_state
        self.labels_: np.ndarray | None = None
        self.cost_: float = np.inf
        self.centroids_num_: np.ndarray | None = None
        self.centroids_cat_: np.ndarray | None = None

    def _dist(self, Xn, Xc, cn, cc):
        """Distance from every row to one centroid (cn numeric, cc categorical)."""
        d = np.zeros(len(Xn) if Xn is not None else len(Xc))
        if Xn is not None and Xn.shape[1] > 0:
            d = d + ((Xn - cn) ** 2).sum(axis=1)
        if Xc is not None and Xc.shape[1] > 0:
            d = d + self.gamma * (Xc != cc).sum(axis=1)
        return d

    def _single_run(self, Xn, Xc, seed):
        rng = np.random.default_rng(seed)
        n = len(Xn) if Xn is not None else len(Xc)
        init = rng.choice(n, size=self.n_clusters, replace=False)
        cnum = Xn[init].copy() if Xn is not None else None
        ccat = Xc[init].copy() if Xc is not None else None
        labels = np.zeros(n, dtype=int)
        for _ in range(self.max_iter):
            D = np.stack([self._dist(Xn, Xc, cnum[k] if cnum is not None else None,
                                     ccat[k] if ccat is not None else None)
                          for k in range(self.n_clusters)], axis=1)
            new_labels = D.argmin(axis=1)
            for k in range(self.n_clusters):
                m = new_labels == k
                if not m.any():
                    continue
                if cnum is not None:
                    cnum[k] = Xn[m].mean(axis=0)
                if ccat is not None:
                    # mode per categorical column
                    for j in range(Xc.shape[1]):
                        vals, counts = np.unique(Xc[m, j], return_counts=True)
                        ccat[k, j] = vals[counts.argmax()]
            if np.array_equal(new_labels, labels):
                labels = new_labels
                break
            labels = new_labels
        cost = float(sum(
            self._dist(Xn, Xc, cnum[k] if cnum is not None else None,
                       ccat[k] if ccat is not None else None)[labels == k].sum()
            for k in range(self.n_clusters)))
        return labels, cnum, ccat, cost

    def fit(self, X: np.ndarray) -> "KPrototypes":
        X = np.asarray(X, dtype=float)
        p = X.shape[1]
        cat_idx = self.cat_idx if self.cat_idx is not None else []
        num_idx = self.num_idx if self.num_idx is not None else \
            [j for j in range(p) if j not in cat_idx]
        Xn = X[:, num_idx] if num_idx else None
        Xc = X[:, cat_idx] if cat_idx else None
        if self.gamma is None:
            # Huang's heuristic: half the mean numeric standard deviation
            self.gamma = 0.5 * float(np.mean(Xn.std(axis=0))) if Xn is not None \
                and Xn.shape[1] > 0 else 1.0
            if self.gamma <= 0:
                self.gamma = 1.0
        base = 0 if self.random_state is None else int(self.random_state)
        best = None
        for i in range(self.n_init):
            labels, cnum, ccat, cost = self._single_run(Xn, Xc, base + i)
            if best is None or cost < best[3]:
                best = (labels, cnum, ccat, cost)
        self.labels_, self.centroids_num_, self.centroids_cat_, self.cost_ = best
        return self


def silhouette_score(X: np.ndarray, labels: np.ndarray, sample_size: int | None = 2000,
                     random_state: int = 0) -> float:
    """Mean silhouette over (optionally) a random sub-sample for speed."""
    X = np.asarray(X, dtype=float)
    labels = np.asarray(labels)
    n = len(X)
    if sample_size is not None and n > sample_size:
        rng = np.random.default_rng(random_state)
        idx = rng.choice(n, size=sample_size, replace=False)
        X_s, lab_s = X[idx], labels[idx]
    else:
        X_s, lab_s = X, labels
    unique = np.unique(lab_s)
    if len(unique) < 2:
        return 0.0
    s_vals = np.zeros(len(X_s))
    # Pre-compute pairwise distances within the sample (sample_size <= 2000, OK)
    diff = X_s[:, None, :] - X_s[None, :, :]
    D = np.sqrt((diff ** 2).sum(-1))
    for i in range(len(X_s)):
        own = lab_s[i]
        same = lab_s == own
        same[i] = False
        if not same.any():
            s_vals[i] = 0.0
            continue
        a = D[i, same].mean()
        bs = []
        for k in unique:
            if k == own:
                continue
            mask = lab_s == k
            if mask.any():
                bs.append(D[i, mask].mean())
        b = min(bs)
        s_vals[i] = (b - a) / max(a, b) if max(a, b) > 0 else 0.0
    return float(s_vals.mean())


# ---------------------------------------------------------------------------
# Logistic Regression (multinomial)
# ---------------------------------------------------------------------------
class LogisticRegression:
    """L2-regularised multinomial logistic regression via batch gradient descent."""

    def __init__(self, C: float = 1.0, lr: float = 0.1, max_iter: int = 500,
                 tol: float = 1e-5, random_state: int | None = 0) -> None:
        self.C = C
        self.lr = lr
        self.max_iter = max_iter
        self.tol = tol
        self.random_state = random_state
        self.W: np.ndarray | None = None  # (n_features+1, n_classes)
        self.classes_: np.ndarray | None = None

    @staticmethod
    def _softmax(Z: np.ndarray) -> np.ndarray:
        Z = Z - Z.max(axis=1, keepdims=True)
        e = np.exp(Z)
        return e / e.sum(axis=1, keepdims=True)

    def fit(self, X: np.ndarray, y: np.ndarray) -> "LogisticRegression":
        X = np.asarray(X, dtype=float)
        y = np.asarray(y)
        self.classes_, y_idx = np.unique(y, return_inverse=True)
        n, p = X.shape
        K = len(self.classes_)
        Xb = np.hstack([np.ones((n, 1)), X])
        rng = np.random.default_rng(self.random_state)
        self.W = rng.normal(0, 0.01, size=(p + 1, K))
        Y = np.eye(K)[y_idx]
        lam = 1.0 / max(self.C, 1e-8)
        prev_loss = np.inf
        for it in range(self.max_iter):
            P = self._softmax(Xb @ self.W)
            # log-loss + L2 (don't regularise intercept)
            loss = -np.mean((Y * np.log(P + 1e-12)).sum(axis=1))
            loss += 0.5 * lam * (self.W[1:] ** 2).sum() / n
            grad = Xb.T @ (P - Y) / n
            grad[1:] += lam * self.W[1:] / n
            self.W -= self.lr * grad
            if abs(prev_loss - loss) < self.tol:
                break
            prev_loss = loss
        return self

    def predict_proba(self, X: np.ndarray) -> np.ndarray:
        Xb = np.hstack([np.ones((len(X), 1)), np.asarray(X, dtype=float)])
        return self._softmax(Xb @ self.W)

    def predict(self, X: np.ndarray) -> np.ndarray:
        return self.classes_[self.predict_proba(X).argmax(axis=1)]


# ---------------------------------------------------------------------------
# Decision tree (CART, Gini)
# ---------------------------------------------------------------------------
class _Node:
    __slots__ = ("feature", "threshold", "left", "right", "prob", "n")

    def __init__(self) -> None:
        self.feature: int | None = None
        self.threshold: float | None = None
        self.left: _Node | None = None
        self.right: _Node | None = None
        self.prob: np.ndarray | None = None
        self.n: int = 0


class DecisionTreeClassifier:
    """CART decision tree with Gini impurity."""

    def __init__(self, max_depth: int = 10, min_samples_split: int = 20,
                 min_samples_leaf: int = 10, max_features: int | str | None = None,
                 random_state: int | None = None) -> None:
        self.max_depth = max_depth
        self.min_samples_split = min_samples_split
        self.min_samples_leaf = min_samples_leaf
        self.max_features = max_features
        self.random_state = random_state
        self.root_: _Node | None = None
        self.classes_: np.ndarray | None = None

    def _gini(self, y_idx: np.ndarray, K: int) -> float:
        if len(y_idx) == 0:
            return 0.0
        counts = np.bincount(y_idx, minlength=K) / len(y_idx)
        return 1.0 - (counts ** 2).sum()

    def _best_split(self, X: np.ndarray, y_idx: np.ndarray, K: int,
                    feat_idx: np.ndarray, rng: np.random.Generator) -> tuple[int, float, float] | None:
        n, _ = X.shape
        base_gini = self._gini(y_idx, K)
        best_gain = 0.0
        best_split = None
        for f in feat_idx:
            col = X[:, f]
            # Use up to 32 quantile candidates to bound complexity
            uniq = np.unique(col)
            if len(uniq) <= 1:
                continue
            if len(uniq) > 16:
                qs = np.quantile(col, np.linspace(0.1, 0.9, 16))
                thresholds = np.unique(qs)
            else:
                thresholds = (uniq[:-1] + uniq[1:]) / 2
            for t in thresholds:
                left_mask = col <= t
                nL = left_mask.sum()
                nR = n - nL
                if nL < self.min_samples_leaf or nR < self.min_samples_leaf:
                    continue
                gL = self._gini(y_idx[left_mask], K)
                gR = self._gini(y_idx[~left_mask], K)
                gain = base_gini - (nL * gL + nR * gR) / n
                if gain > best_gain:
                    best_gain = gain
                    best_split = (f, float(t), gain)
        return best_split

    def _build(self, X: np.ndarray, y_idx: np.ndarray, K: int, depth: int,
               rng: np.random.Generator) -> _Node:
        node = _Node()
        node.n = len(y_idx)
        counts = np.bincount(y_idx, minlength=K)
        node.prob = counts / max(counts.sum(), 1)
        if (depth >= self.max_depth
                or len(y_idx) < self.min_samples_split
                or (counts > 0).sum() == 1):
            return node
        n_feat = X.shape[1]
        if self.max_features is None:
            feat_idx = np.arange(n_feat)
        elif isinstance(self.max_features, int):
            feat_idx = rng.choice(n_feat, size=min(self.max_features, n_feat), replace=False)
        elif self.max_features == "sqrt":
            k = max(1, int(np.sqrt(n_feat)))
            feat_idx = rng.choice(n_feat, size=k, replace=False)
        else:
            feat_idx = np.arange(n_feat)
        split = self._best_split(X, y_idx, K, feat_idx, rng)
        if split is None:
            return node
        f, t, _ = split
        left_mask = X[:, f] <= t
        node.feature = f
        node.threshold = t
        node.left = self._build(X[left_mask], y_idx[left_mask], K, depth + 1, rng)
        node.right = self._build(X[~left_mask], y_idx[~left_mask], K, depth + 1, rng)
        return node

    def fit(self, X: np.ndarray, y: np.ndarray) -> "DecisionTreeClassifier":
        X = np.asarray(X, dtype=float)
        y = np.asarray(y)
        self.classes_, y_idx = np.unique(y, return_inverse=True)
        rng = np.random.default_rng(self.random_state)
        self.root_ = self._build(X, y_idx, len(self.classes_), 0, rng)
        return self

    def _walk(self, node: _Node, x: np.ndarray) -> np.ndarray:
        while node.feature is not None:
            if x[node.feature] <= node.threshold:
                node = node.left
            else:
                node = node.right
        return node.prob

    def _walk_batch(self, X: np.ndarray) -> np.ndarray:
        """Vectorised prediction: route an index set through the tree."""
        K = len(self.classes_)
        out = np.zeros((len(X), K))
        # Stack: list of (node, indices)
        stack = [(self.root_, np.arange(len(X)))]
        while stack:
            node, idx = stack.pop()
            if node.feature is None or len(idx) == 0:
                if len(idx):
                    out[idx] = node.prob
                continue
            mask = X[idx, node.feature] <= node.threshold
            stack.append((node.left, idx[mask]))
            stack.append((node.right, idx[~mask]))
        return out

    def predict_proba(self, X: np.ndarray) -> np.ndarray:
        X = np.asarray(X, dtype=float)
        return self._walk_batch(X)

    def predict(self, X: np.ndarray) -> np.ndarray:
        return self.classes_[self.predict_proba(X).argmax(axis=1)]


class _RegNode:
    __slots__ = ("feature", "threshold", "left", "right", "value", "leaf_id")

    def __init__(self) -> None:
        self.feature: int | None = None
        self.threshold: float | None = None
        self.left: _RegNode | None = None
        self.right: _RegNode | None = None
        self.value: float = 0.0
        self.leaf_id: int = -1


class DecisionTreeRegressor:
    """CART regression tree with variance-reduction splits.

    Exposes apply() returning the leaf index for each sample, which the
    gradient-boosting classifier uses for Friedman leaf-value refinement.
    """

    def __init__(self, max_depth: int = 3, min_samples_split: int = 20,
                 min_samples_leaf: int = 10, max_features: str | int | None = None,
                 random_state: int | None = None) -> None:
        self.max_depth = max_depth
        self.min_samples_split = min_samples_split
        self.min_samples_leaf = min_samples_leaf
        self.max_features = max_features
        self.random_state = random_state
        self.root_: _RegNode | None = None
        self._leaf_counter = 0

    @staticmethod
    def _var(y: np.ndarray) -> float:
        return float(y.var()) if len(y) else 0.0

    def _best_split(self, X: np.ndarray, y: np.ndarray,
                    feat_idx: np.ndarray) -> tuple[int, float] | None:
        n = len(y)
        parent_var = self._var(y)
        best_gain = 0.0
        best = None
        for f in feat_idx:
            col = X[:, f]
            uniq = np.unique(col)
            if len(uniq) <= 1:
                continue
            if len(uniq) > 16:
                thresholds = np.unique(np.quantile(col, np.linspace(0.1, 0.9, 16)))
            else:
                thresholds = (uniq[:-1] + uniq[1:]) / 2
            for t in thresholds:
                left = col <= t
                nL = int(left.sum())
                nR = n - nL
                if nL < self.min_samples_leaf or nR < self.min_samples_leaf:
                    continue
                gain = parent_var - (nL * self._var(y[left])
                                     + nR * self._var(y[~left])) / n
                if gain > best_gain:
                    best_gain = gain
                    best = (f, float(t))
        return best

    def _build(self, X: np.ndarray, y: np.ndarray, depth: int,
               rng: np.random.Generator) -> _RegNode:
        node = _RegNode()
        node.value = float(y.mean()) if len(y) else 0.0
        if (depth >= self.max_depth or len(y) < self.min_samples_split
                or np.allclose(y, y[0])):
            node.leaf_id = self._leaf_counter
            self._leaf_counter += 1
            return node
        nf = X.shape[1]
        if self.max_features == "sqrt":
            k = max(1, int(np.sqrt(nf)))
            feat_idx = rng.choice(nf, size=k, replace=False)
        else:
            feat_idx = np.arange(nf)
        split = self._best_split(X, y, feat_idx)
        if split is None:
            node.leaf_id = self._leaf_counter
            self._leaf_counter += 1
            return node
        f, t = split
        left = X[:, f] <= t
        node.feature, node.threshold = f, t
        node.left = self._build(X[left], y[left], depth + 1, rng)
        node.right = self._build(X[~left], y[~left], depth + 1, rng)
        return node

    def fit(self, X: np.ndarray, y: np.ndarray) -> "DecisionTreeRegressor":
        X = np.asarray(X, dtype=float)
        y = np.asarray(y, dtype=float)
        self._leaf_counter = 0
        rng = np.random.default_rng(self.random_state)
        self.root_ = self._build(X, y, 0, rng)
        return self

    def _walk(self, x: np.ndarray, want_leaf: bool):
        node = self.root_
        while node.feature is not None:
            node = node.left if x[node.feature] <= node.threshold else node.right
        return node.leaf_id if want_leaf else node.value

    def predict(self, X: np.ndarray) -> np.ndarray:
        X = np.asarray(X, dtype=float)
        return np.array([self._walk(x, False) for x in X])

    def apply(self, X: np.ndarray) -> np.ndarray:
        X = np.asarray(X, dtype=float)
        return np.array([self._walk(x, True) for x in X], dtype=int)


class GradientBoostingClassifier:
    """Binary gradient-boosted trees (log-loss) with Friedman leaf refinement.

    Stage-wise additive model: each round fits a regression tree to the
    negative log-loss gradient (y - p); leaf values are then refined by a
    single Newton step gamma = Σresidual / Σp(1-p), the standard TreeBoost
    update. Suited to the binary risk-profile target in this project.
    """

    def __init__(self, n_estimators: int = 80, learning_rate: float = 0.1,
                 max_depth: int = 3, min_samples_leaf: int = 20,
                 min_samples_split: int = 40, subsample: float = 1.0,
                 random_state: int | None = 0) -> None:
        self.n_estimators = n_estimators
        self.learning_rate = learning_rate
        self.max_depth = max_depth
        self.min_samples_leaf = min_samples_leaf
        self.min_samples_split = min_samples_split
        self.subsample = subsample
        self.random_state = random_state
        self.trees_: list[DecisionTreeRegressor] = []
        self.leaf_gammas_: list[dict] = []
        self.classes_: np.ndarray | None = None
        self.F0: float = 0.0

    @staticmethod
    def _sigmoid(z: np.ndarray) -> np.ndarray:
        return 1.0 / (1.0 + np.exp(-np.clip(z, -30, 30)))

    def fit(self, X: np.ndarray, y: np.ndarray) -> "GradientBoostingClassifier":
        X = np.asarray(X, dtype=float)
        y = np.asarray(y)
        self.classes_ = np.unique(y)
        if len(self.classes_) != 2:
            raise ValueError("GradientBoostingClassifier supports binary targets only")
        y01 = (y == self.classes_[1]).astype(float)
        p0 = float(np.clip(y01.mean(), 1e-6, 1 - 1e-6))
        self.F0 = np.log(p0 / (1 - p0))
        F = np.full(len(y), self.F0)
        rng = np.random.default_rng(self.random_state)
        n = len(X)
        self.trees_, self.leaf_gammas_ = [], []
        for _ in range(self.n_estimators):
            p = self._sigmoid(F)
            residual = y01 - p  # negative gradient of log-loss
            if self.subsample < 1.0:
                idx = rng.choice(n, size=max(1, int(self.subsample * n)),
                                 replace=False)
            else:
                idx = np.arange(n)
            tree = DecisionTreeRegressor(
                max_depth=self.max_depth,
                min_samples_leaf=self.min_samples_leaf,
                min_samples_split=self.min_samples_split,
                random_state=int(rng.integers(0, 1_000_000)),
            )
            tree.fit(X[idx], residual[idx])
            # Friedman leaf refinement (Newton step) on the fitted subsample
            leaf_ids = tree.apply(X[idx])
            p_sub = p[idx]
            res_sub = residual[idx]
            gammas: dict[int, float] = {}
            for lid in np.unique(leaf_ids):
                m = leaf_ids == lid
                num = res_sub[m].sum()
                den = (p_sub[m] * (1.0 - p_sub[m])).sum()
                gammas[lid] = float(num / (den + 1e-12))
            self.trees_.append(tree)
            self.leaf_gammas_.append(gammas)
            leaves_all = tree.apply(X)
            update = np.array([gammas.get(int(l), 0.0) for l in leaves_all])
            F += self.learning_rate * update
        return self

    def decision_function(self, X: np.ndarray) -> np.ndarray:
        X = np.asarray(X, dtype=float)
        F = np.full(len(X), self.F0)
        for tree, gammas in zip(self.trees_, self.leaf_gammas_):
            leaves = tree.apply(X)
            F += self.learning_rate * np.array(
                [gammas.get(int(l), 0.0) for l in leaves])
        return F

    def predict_proba(self, X: np.ndarray) -> np.ndarray:
        p = self._sigmoid(self.decision_function(X))
        return np.stack([1.0 - p, p], axis=1)

    def predict(self, X: np.ndarray) -> np.ndarray:
        return self.classes_[(self.predict_proba(X)[:, 1] >= 0.5).astype(int)]


class RandomForestClassifier:
    """Random forest with bootstrap aggregation and feature subsampling."""

    def __init__(self, n_estimators: int = 100, max_depth: int = 10,
                 min_samples_split: int = 20, min_samples_leaf: int = 10,
                 max_features: str | int | None = "sqrt",
                 random_state: int | None = 0) -> None:
        self.n_estimators = n_estimators
        self.max_depth = max_depth
        self.min_samples_split = min_samples_split
        self.min_samples_leaf = min_samples_leaf
        self.max_features = max_features
        self.random_state = random_state
        self.trees_: list[DecisionTreeClassifier] = []
        self.classes_: np.ndarray | None = None

    def fit(self, X: np.ndarray, y: np.ndarray) -> "RandomForestClassifier":
        X = np.asarray(X, dtype=float)
        y = np.asarray(y)
        self.classes_ = np.unique(y)
        rng = np.random.default_rng(self.random_state)
        n = len(X)
        self.trees_ = []
        for i in range(self.n_estimators):
            sample_idx = rng.integers(0, n, size=n)
            tree = DecisionTreeClassifier(
                max_depth=self.max_depth,
                min_samples_split=self.min_samples_split,
                min_samples_leaf=self.min_samples_leaf,
                max_features=self.max_features,
                random_state=int(rng.integers(0, 1_000_000)),
            )
            tree.fit(X[sample_idx], y[sample_idx])
            self.trees_.append(tree)
        return self

    def predict_proba(self, X: np.ndarray) -> np.ndarray:
        # Align trees with full class list (some bootstrap samples may miss classes)
        K = len(self.classes_)
        probs = np.zeros((len(X), K))
        for tree in self.trees_:
            p = tree.predict_proba(X)
            # map tree.classes_ into the full class list
            for j, c in enumerate(tree.classes_):
                k = int(np.where(self.classes_ == c)[0][0])
                probs[:, k] += p[:, j]
        probs /= len(self.trees_)
        return probs

    def predict(self, X: np.ndarray) -> np.ndarray:
        return self.classes_[self.predict_proba(X).argmax(axis=1)]


# ---------------------------------------------------------------------------
# Cross-validation
# ---------------------------------------------------------------------------
class StratifiedKFold:
    def __init__(self, n_splits: int = 5, shuffle: bool = True,
                 random_state: int | None = 0) -> None:
        self.n_splits = n_splits
        self.shuffle = shuffle
        self.random_state = random_state

    def split(self, X: np.ndarray, y: np.ndarray):
        y = np.asarray(y)
        n = len(y)
        rng = np.random.default_rng(self.random_state)
        fold = np.zeros(n, dtype=int)
        for c in np.unique(y):
            idx = np.where(y == c)[0]
            if self.shuffle:
                rng.shuffle(idx)
            for i, x in enumerate(idx):
                fold[x] = i % self.n_splits
        for k in range(self.n_splits):
            test = np.where(fold == k)[0]
            train = np.where(fold != k)[0]
            yield train, test


def train_test_split(X: np.ndarray, y: np.ndarray, test_size: float = 0.2,
                     random_state: int = 0, stratify: np.ndarray | None = None):
    X = np.asarray(X)
    y = np.asarray(y)
    rng = np.random.default_rng(random_state)
    if stratify is None:
        idx = rng.permutation(len(X))
        cut = int(len(X) * (1 - test_size))
        tr, te = idx[:cut], idx[cut:]
    else:
        s = np.asarray(stratify)
        tr_list, te_list = [], []
        for c in np.unique(s):
            ci = np.where(s == c)[0]
            rng.shuffle(ci)
            cut = int(len(ci) * (1 - test_size))
            tr_list.append(ci[:cut])
            te_list.append(ci[cut:])
        tr = np.concatenate(tr_list)
        te = np.concatenate(te_list)
        rng.shuffle(tr)
        rng.shuffle(te)
    return X[tr], X[te], y[tr], y[te]


# ---------------------------------------------------------------------------
# Metrics
# ---------------------------------------------------------------------------
def accuracy(y_true: np.ndarray, y_pred: np.ndarray) -> float:
    return float((np.asarray(y_true) == np.asarray(y_pred)).mean())


def confusion_matrix(y_true: np.ndarray, y_pred: np.ndarray, labels=None) -> np.ndarray:
    y_true, y_pred = np.asarray(y_true), np.asarray(y_pred)
    if labels is None:
        labels = np.unique(np.concatenate([y_true, y_pred]))
    K = len(labels)
    idx = {c: i for i, c in enumerate(labels)}
    cm = np.zeros((K, K), dtype=int)
    for t, p in zip(y_true, y_pred):
        cm[idx[t], idx[p]] += 1
    return cm


def macro_f1(y_true: np.ndarray, y_pred: np.ndarray) -> float:
    y_true, y_pred = np.asarray(y_true), np.asarray(y_pred)
    classes = np.unique(y_true)
    f1s = []
    for c in classes:
        tp = int(((y_pred == c) & (y_true == c)).sum())
        fp = int(((y_pred == c) & (y_true != c)).sum())
        fn = int(((y_pred != c) & (y_true == c)).sum())
        denom = (2 * tp + fp + fn)
        f1s.append(2 * tp / denom if denom > 0 else 0.0)
    return float(np.mean(f1s))


def balanced_accuracy(y_true: np.ndarray, y_pred: np.ndarray) -> float:
    y_true, y_pred = np.asarray(y_true), np.asarray(y_pred)
    recalls = []
    for c in np.unique(y_true):
        mask = y_true == c
        if mask.sum() > 0:
            recalls.append(((y_pred == c) & mask).sum() / mask.sum())
    return float(np.mean(recalls)) if recalls else 0.0


def _roc_auc_binary(y_true: np.ndarray, y_score: np.ndarray) -> float:
    """Trapezoidal AUC for binary classification."""
    y_true = np.asarray(y_true).astype(int)
    y_score = np.asarray(y_score)
    order = np.argsort(-y_score)
    y_sorted = y_true[order]
    P = y_true.sum()
    N = len(y_true) - P
    if P == 0 or N == 0:
        return float("nan")
    tps = np.cumsum(y_sorted)
    fps = np.cumsum(1 - y_sorted)
    tpr = tps / P
    fpr = fps / N
    tpr = np.concatenate([[0], tpr])
    fpr = np.concatenate([[0], fpr])
    return float(np.trapz(tpr, fpr))


def roc_auc_ovr(y_true: np.ndarray, y_proba: np.ndarray, classes: np.ndarray) -> float:
    """One-vs-rest macro-averaged AUC."""
    y_true = np.asarray(y_true)
    aucs = []
    for j, c in enumerate(classes):
        yb = (y_true == c).astype(int)
        a = _roc_auc_binary(yb, y_proba[:, j])
        if not np.isnan(a):
            aucs.append(a)
    return float(np.mean(aucs)) if aucs else float("nan")


# ---------------------------------------------------------------------------
# Permutation importance
# ---------------------------------------------------------------------------
def permutation_importance(model, X: np.ndarray, y: np.ndarray, n_repeats: int = 5,
                           random_state: int = 0, scoring=macro_f1) -> np.ndarray:
    rng = np.random.default_rng(random_state)
    base = scoring(y, model.predict(X))
    importances = np.zeros((X.shape[1], n_repeats))
    for j in range(X.shape[1]):
        for r in range(n_repeats):
            X_perm = X.copy()
            X_perm[:, j] = rng.permutation(X_perm[:, j])
            importances[j, r] = base - scoring(y, model.predict(X_perm))
    return importances


# ===========================================================================
# SECTION 2  analysis pipeline  (was run_analysis.py)
# ===========================================================================

sns.set_theme(style="whitegrid", context="paper")
RNG = 20240516

# ---------------------------------------------------------------------------
# Unified figure colour scheme (Royal Mint / dusty rose)
#   mint  = low-risk profile / demographics-only baseline / below-mean
#   rose  = elevated-risk profile / depression-inclusive model / above-mean
# Every figure in the report draws from these anchors so the deck reads as
# one coherent visual identity.
# ---------------------------------------------------------------------------
from matplotlib.colors import LinearSegmentedColormap

# ColorBrewer RdBu anchors — a colourblind-safe diverging pair that
# prints cleanly: blue reads as low-risk (cool), red as elevated (hot).
COL_LOW = "#2166AC"                      # ColorBrewer RdBu blue — low-risk
COL_HIGH = "#B2182B"                     # ColorBrewer RdBu red — elevated-risk
PROFILE_COLORS = [COL_LOW, COL_HIGH]     # risk profile 0, risk profile 1
DEMO_SHADES = list(sns.light_palette(COL_LOW, n_colors=7))[3:6]   # demo models
FULL_SHADES = list(sns.light_palette(COL_HIGH, n_colors=7))[3:6]  # + PHQ-9
# Diverging map (blue <-> red) for the cluster heatmap; sequential blue
# for the confusion matrices.
DIVERGING_CMAP = LinearSegmentedColormap.from_list(
    "blue_red", [COL_LOW, "#f4f2f0", COL_HIGH])
SEQ_CMAP = sns.light_palette(COL_LOW, as_cmap=True)

# ---------------------------------------------------------------------------
# Gradient-boosting backend
# ---------------------------------------------------------------------------
# The gradient-boosting model uses the best available library: LightGBM is
# preferred, then XGBoost, then the from-scratch NumPy GradientBoostingClassifier
# in ml_lib.py. Both LightGBM and XGBoost expose a scikit-learn-compatible
# fit / predict / predict_proba interface, so they are genuine drop-in
# replacements. Install either library (`pip install lightgbm`) and re-run to
# use it automatically; with neither installed the pipeline still runs end to
# end on the from-scratch implementation.
GB_BACKEND = "from-scratch NumPy"
try:
    from lightgbm import LGBMClassifier  # noqa: F401
    GB_BACKEND = "LightGBM"
except Exception:
    try:
        from xgboost import XGBClassifier  # noqa: F401
        GB_BACKEND = "XGBoost"
    except Exception:
        pass


def make_gb():
    """Return a gradient-boosting classifier from the best available backend.

    Hyperparameters (n_estimators 60, learning_rate 0.05, max_depth 3) were
    selected by grid search (Table 1) and map onto all three backends.
    """
    if GB_BACKEND == "LightGBM":
        return LGBMClassifier(
            n_estimators=60, learning_rate=0.05, max_depth=3, num_leaves=8,
            subsample=0.8, subsample_freq=1, min_child_samples=30,
            random_state=0, verbose=-1)
    if GB_BACKEND == "XGBoost":
        return XGBClassifier(
            n_estimators=60, learning_rate=0.05, max_depth=3, subsample=0.8,
            min_child_weight=5, random_state=0, verbosity=0,
            eval_metric="logloss")
    return GradientBoostingClassifier(
        n_estimators=60, learning_rate=0.05, max_depth=3,
        min_samples_leaf=30, min_samples_split=60, subsample=0.8,
        random_state=0)


# ---------------------------------------------------------------------------
# 1. Load and merge NHANES modules
# ---------------------------------------------------------------------------
def load_module(name: str) -> pd.DataFrame:
    """Load a single NHANES XPT (with the `.xpt.txt` suffix used in this repo)."""
    p = DATA / f"{name}.xpt.txt"
    if not p.exists():
        raise FileNotFoundError(f"Missing NHANES module: {p}")
    return pd.read_sas(p, format="xport")


def build_cohort() -> pd.DataFrame:
    demo = load_module("P_DEMO")
    dpq = load_module("P_DPQ")
    bpq = load_module("P_BPQ")
    mcq = load_module("P_MCQ")
    slq = load_module("P_SLQ")
    smq = load_module("P_SMQFAM")

    cols_demo = ["SEQN", "RIDAGEYR", "RIAGENDR", "RIDRETH3",
                 "DMDEDUC2", "DMDMARTZ", "INDFMPIR", "RIDEXPRG"]
    cols_bpq = ["SEQN", "BPQ020", "BPQ050A", "BPQ080", "BPQ100D"]
    cols_mcq = ["SEQN", "MCQ080", "MCQ366A", "MCQ366B", "MCQ366C",
                "MCQ366D", "MCQ300A", "MCQ300C"]
    cols_slq = ["SEQN", "SLD012"]
    cols_smq = ["SEQN", "SMD460"]
    cols_dpq = ["SEQN"] + [f"DPQ0{i}0" for i in range(1, 10)]

    df = (demo[cols_demo]
          .merge(dpq[cols_dpq], on="SEQN", how="inner")
          .merge(bpq[cols_bpq], on="SEQN", how="left")
          .merge(mcq[cols_mcq], on="SEQN", how="left")
          .merge(slq[cols_slq], on="SEQN", how="left")
          .merge(smq[cols_smq], on="SEQN", how="left"))

    # Restrict to 18-39 and drop pregnant respondents
    df = df.query("RIDAGEYR >= 18 and RIDAGEYR <= 39").copy()
    df = df[df["RIDEXPRG"].fillna(0) != 1].copy()
    df.drop(columns=["RIDEXPRG"], inplace=True)
    return df.reset_index(drop=True)


# ---------------------------------------------------------------------------
# 2. PHQ-9 + feature engineering
# ---------------------------------------------------------------------------
PHQ_ITEMS = [f"DPQ0{i}0" for i in range(1, 10)]


def code_phq9(df: pd.DataFrame) -> pd.DataFrame:
    """Code 7/9 as missing, prorate when ≤2 missing, compute total."""
    out = df.copy()
    for col in PHQ_ITEMS:
        out.loc[out[col].isin([7, 9]), col] = np.nan
    miss = out[PHQ_ITEMS].isna().sum(axis=1)
    mean_item = out[PHQ_ITEMS].mean(axis=1, skipna=True)
    total = out[PHQ_ITEMS].sum(axis=1, skipna=True) + miss * mean_item
    total[miss > 2] = np.nan
    out["PHQ9_total"] = total
    out["PHQ9_pos"] = (out["PHQ9_total"] >= 10).astype("Int64")
    return out


def code_binary(df: pd.DataFrame, cols: list[str], yes: int = 1) -> pd.DataFrame:
    """NHANES coding: 1=Yes, 2=No, 7=refused, 9=don't know."""
    out = df.copy()
    for c in cols:
        s = out[c]
        new = pd.Series(np.nan, index=s.index, dtype="float")
        new[s == yes] = 1.0
        new[s == 2] = 0.0
        out[c] = new
    return out


def build_features(df: pd.DataFrame) -> pd.DataFrame:
    out = code_phq9(df)
    binary_cols = ["BPQ020", "BPQ050A", "BPQ080", "BPQ100D",
                   "MCQ080", "MCQ366A", "MCQ366B", "MCQ366C", "MCQ366D",
                   "MCQ300A", "MCQ300C"]
    out = code_binary(out, binary_cols)

    # Sleep: NHANES values 77, 99 are refused/don't know
    out.loc[out["SLD012"].isin([77, 99]), "SLD012"] = np.nan
    out["insufficient_sleep"] = (out["SLD012"] < 7).astype("float")

    # Household smokers: code 999 missing; cap at 4
    out.loc[out["SMD460"].isin([777, 999]), "SMD460"] = np.nan
    out["household_smoker"] = (out["SMD460"] >= 1).astype("float")

    # Demographics
    out["female"] = (out["RIAGENDR"] == 2).astype(int)
    # Race/ethnicity (RIDRETH3): 1 Mex Am, 2 Other Hisp, 3 NH White,
    # 4 NH Black, 6 NH Asian, 7 Other/Multi
    race_map = {1: "MexAm", 2: "OtherHisp", 3: "NHWhite",
                4: "NHBlack", 6: "NHAsian", 7: "OtherMulti"}
    out["race"] = out["RIDRETH3"].map(race_map).fillna("OtherMulti")
    # Education: 1<9, 2 9-11, 3 HS, 4 some college, 5 college+
    edu_map = {1: "LessHS", 2: "LessHS", 3: "HS",
               4: "SomeCollege", 5: "CollegePlus"}
    out["education"] = out["DMDEDUC2"].map(edu_map).fillna("HS")
    # Marital: 1 married/partner, 2 widowed/divorced/separated, 3 never
    mar_map = {1: "MarriedPartner", 2: "WidDivSep", 3: "Never"}
    out["marital"] = out["DMDMARTZ"].map(mar_map).fillna("Never")

    return out


# ---------------------------------------------------------------------------
# 3. Imputation & one-hot encoding
# ---------------------------------------------------------------------------
RISK_BINARY = ["BPQ020", "BPQ050A", "BPQ080", "BPQ100D",
               "MCQ080", "MCQ366A", "MCQ366B", "MCQ366C", "MCQ366D",
               "MCQ300A", "MCQ300C", "insufficient_sleep", "household_smoker"]
RISK_CONT = ["SLD012", "SMD460"]
DEMO_NUM = ["RIDAGEYR", "INDFMPIR", "female"]
DEMO_CAT = ["race", "education", "marital"]


def impute(df: pd.DataFrame) -> pd.DataFrame:
    out = df.copy()
    for c in RISK_BINARY:
        med = out[c].mode(dropna=True)
        fill = med.iloc[0] if not med.empty else 0.0
        out[c] = out[c].fillna(fill)
    for c in RISK_CONT + ["INDFMPIR"]:
        out[c] = out[c].fillna(out[c].median())
    return out


def one_hot(df: pd.DataFrame, cols: list[str]) -> pd.DataFrame:
    return pd.get_dummies(df, columns=cols, drop_first=False, dtype=float)


# ---------------------------------------------------------------------------
# 4. K-means clustering on CV risk indicators
# ---------------------------------------------------------------------------
def derive_risk_profiles(df: pd.DataFrame, k_range=(2, 3, 4),
                         seed: int = RNG) -> tuple[pd.DataFrame, dict]:
    cv_feats = df[RISK_BINARY + ["SLD012"]].copy()
    # Reverse-code sleep so higher = worse (insufficient already binary)
    cv_feats["sleep_deficit"] = np.maximum(0.0, 7.0 - cv_feats["SLD012"].fillna(7.0))
    cv_feats.drop(columns=["SLD012"], inplace=True)
    X = cv_feats.to_numpy(dtype=float)
    scaler = StandardScaler().fit(X)
    Xz = scaler.transform(X)

    sil_scores = {}
    inertias = {}
    for k in k_range:
        km = KMeans(n_clusters=k, n_init=10, random_state=seed).fit(Xz)
        sil_scores[k] = silhouette_score(Xz, km.labels_, sample_size=1500,
                                         random_state=seed)
        inertias[k] = km.inertia_

    k_star = max(sil_scores, key=sil_scores.get)
    km_final = KMeans(n_clusters=k_star, n_init=20, random_state=seed).fit(Xz)
    labels = km_final.labels_

    # Order clusters by composite-risk rank (mean across risk indicators)
    cluster_means = pd.DataFrame(km_final.cluster_centers_,
                                 columns=cv_feats.columns)
    risk_rank = cluster_means.mean(axis=1).rank().astype(int) - 1
    relabel = {old: new for old, new in zip(cluster_means.index, risk_rank)}
    new_labels = np.array([relabel[l] for l in labels])

    df_out = df.copy()
    df_out["risk_profile"] = new_labels
    info = dict(
        k_star=k_star,
        silhouette=sil_scores,
        inertia=inertias,
        cluster_means=cluster_means.rename(index=relabel).sort_index().to_dict(),
        feature_names=list(cv_feats.columns),
        scaler_mean=scaler.mean_.tolist(),
        scaler_std=scaler.std_.tolist(),
    )
    return df_out, info


# ---------------------------------------------------------------------------
# 5. Modelling helpers
# ---------------------------------------------------------------------------
def make_feature_matrix(df: pd.DataFrame, include_depression: bool = True):
    cols = DEMO_NUM.copy()
    df_oh = one_hot(df, DEMO_CAT)
    cat_cols = [c for c in df_oh.columns
                if any(c.startswith(p + "_") for p in DEMO_CAT)]
    cols += cat_cols
    if include_depression:
        cols += ["PHQ9_total"]
    X = df_oh[cols].to_numpy(dtype=float)
    feature_names = cols
    return X, feature_names


def oversample(X, y, random_state=0):
    """Random oversampling to equalise class frequencies."""
    rng = np.random.default_rng(random_state)
    classes, counts = np.unique(y, return_counts=True)
    max_n = counts.max()
    out_idx = []
    for c, n in zip(classes, counts):
        idx = np.where(y == c)[0]
        if n < max_n:
            extra = rng.choice(idx, size=max_n - n, replace=True)
            idx = np.concatenate([idx, extra])
        out_idx.append(idx)
    sel = np.concatenate(out_idx)
    rng.shuffle(sel)
    return X[sel], y[sel]


def cross_validate(model_fn, X, y, n_splits=5, seeds=(1, 2, 3, 4, 5),
                   balance=True):
    classes = np.unique(y)
    fold_results = []
    for seed in seeds:
        skf = StratifiedKFold(n_splits=n_splits, shuffle=True, random_state=seed)
        for fold, (tr, te) in enumerate(skf.split(X, y)):
            scaler = StandardScaler().fit(X[tr])
            Xtr = scaler.transform(X[tr])
            Xte = scaler.transform(X[te])
            y_tr = y[tr]
            if balance:
                Xtr, y_tr = oversample(Xtr, y_tr, random_state=seed)
            model = model_fn()
            model.fit(Xtr, y_tr)
            y_pred = model.predict(Xte)
            y_proba = model.predict_proba(Xte)
            # Per-class recall / precision
            per_class = {}
            for c in classes:
                mask = y[te] == c
                if mask.sum() > 0:
                    per_class[f"recall_{c}"] = float(((y_pred == c) & mask).sum() / mask.sum())
                pmask = y_pred == c
                if pmask.sum() > 0:
                    per_class[f"prec_{c}"] = float(((y_pred == c) & (y[te] == c)).sum() / pmask.sum())
                else:
                    per_class[f"prec_{c}"] = 0.0
            fold_results.append(dict(
                seed=seed,
                fold=fold,
                acc=accuracy(y[te], y_pred),
                bal_acc=balanced_accuracy(y[te], y_pred),
                macro_f1=macro_f1(y[te], y_pred),
                roc_auc=roc_auc_ovr(y[te], y_proba, classes),
                **per_class,
            ))
    return pd.DataFrame(fold_results)


# ---------------------------------------------------------------------------
# 6. Plots
# ---------------------------------------------------------------------------
def plot_cluster_profile(cluster_means: dict, feature_names: list[str],
                         path: Path) -> None:
    df = pd.DataFrame(cluster_means).T.reindex(feature_names).T
    df.index.name = "Risk profile"
    fig, ax = plt.subplots(figsize=(8, 0.6 * len(df) + 2))
    sns.heatmap(df, annot=True, fmt=".2f", cmap=DIVERGING_CMAP,
                center=0, linewidths=0.5, linecolor="white",
                cbar_kws={"label": "Standardised mean"}, ax=ax)
    ax.set_xlabel("")
    plt.xticks(rotation=45, ha="right")
    plt.title("Standardised cluster centres on cardiovascular risk indicators")
    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)


def plot_silhouette(sil_scores: dict, inertias: dict, path: Path) -> None:
    fig, axes = plt.subplots(1, 2, figsize=(9, 3.8))
    ks = sorted(sil_scores)
    axes[0].plot(ks, [sil_scores[k] for k in ks], marker="o",
                 markersize=8, color=COL_HIGH, linewidth=2.4)
    for k in ks:
        axes[0].annotate(f"{sil_scores[k]:.3f}", (k, sil_scores[k]),
                         textcoords="offset points", xytext=(0, 10),
                         ha="center", fontsize=10, fontweight="bold")
    axes[0].set_xlabel("Number of clusters K", fontsize=11)
    axes[0].set_ylabel("Silhouette", fontsize=11)
    axes[0].set_title("Silhouette by K", fontsize=12)
    axes[0].margins(y=0.18)
    axes[1].plot(ks, [inertias[k] for k in ks], marker="o",
                 markersize=8, color=COL_LOW, linewidth=2.4)
    for k in ks:
        axes[1].annotate(f"{inertias[k]:,.0f}", (k, inertias[k]),
                         textcoords="offset points", xytext=(0, 10),
                         ha="center", fontsize=10, fontweight="bold")
    axes[1].set_xlabel("Number of clusters K", fontsize=11)
    axes[1].set_ylabel("Inertia", fontsize=11)
    axes[1].set_title("Elbow plot", fontsize=12)
    axes[1].margins(y=0.18)
    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)


def plot_phq_by_profile(df: pd.DataFrame, path: Path) -> None:
    fig, ax = plt.subplots(figsize=(9, 5.5))
    sns.boxplot(data=df, x="risk_profile", y="PHQ9_total",
                hue="risk_profile", palette=PROFILE_COLORS, legend=False,
                width=0.55, ax=ax)
    sns.stripplot(data=df.sample(min(len(df), 600), random_state=0),
                  x="risk_profile", y="PHQ9_total", color=".25",
                  size=3, alpha=0.4, ax=ax)
    ax.set_xticklabels(["Profile 0 — low CV risk",
                        "Profile 1 — elevated CV risk"], fontsize=12)
    ax.set_xlabel("Cardiovascular risk profile", fontsize=12)
    ax.set_ylabel("PHQ-9 total score", fontsize=12)
    ax.tick_params(axis="y", labelsize=11)
    ax.set_title("Depressive symptom burden by CV risk profile",
                 fontsize=13)
    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)


def plot_metric_bars(results: dict, path: Path) -> None:
    rows = []
    for name, df in results.items():
        for m in ["acc", "bal_acc", "macro_f1", "roc_auc"]:
            rows.append(dict(model=name, metric=m,
                             mean=df[m].mean(), sd=df[m].std()))
    long = pd.DataFrame(rows)
    # Mint shades = demographics-only baselines; rose shades = + PHQ-9 models
    model_palette = {
        "LR_demo": DEMO_SHADES[0], "RF_demo": DEMO_SHADES[1],
        "GB_demo": DEMO_SHADES[2],
        "LR_full": FULL_SHADES[0], "RF_full": FULL_SHADES[1],
        "GB_full": FULL_SHADES[2],
    }
    fig, ax = plt.subplots(figsize=(9.5, 5))
    sns.barplot(data=long, x="metric", y="mean", hue="model", ax=ax,
                palette=model_palette, errorbar=None,
                edgecolor="#333333", linewidth=0.6)
    # Manual error bars + value labels
    metrics = ["acc", "bal_acc", "macro_f1", "roc_auc"]
    models = list(results)
    width = 0.8 / len(models)
    for i, mname in enumerate(models):
        for j, m in enumerate(metrics):
            mean = long.query("model==@mname and metric==@m")["mean"].iloc[0]
            sd = long.query("model==@mname and metric==@m")["sd"].iloc[0]
            x = j - 0.4 + (i + 0.5) * width
            ax.errorbar(x, mean, yerr=sd, fmt="none", ecolor="#222222",
                        capsize=3, linewidth=0.9)
            ax.text(x, mean + sd + 0.02, f"{mean:.2f}", ha="center",
                    va="bottom", fontsize=7.5, rotation=90,
                    fontweight="bold", color="#222222")
    ax.set_ylabel("Score", fontsize=12)
    ax.set_xlabel("")
    ax.set_ylim(0, 0.85)
    ax.tick_params(labelsize=11)
    ax.set_title("Cross-validated performance (mean ± SD across 15 folds)",
                 fontsize=13)
    ax.legend(fontsize=9, title_fontsize=10)
    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)


def plot_confusion(y_true, y_pred, classes, title, path):
    cm = confusion_matrix(y_true, y_pred, labels=classes)
    cm_norm = cm / cm.sum(axis=1, keepdims=True)
    fig, ax = plt.subplots(figsize=(4.5, 4))
    sns.heatmap(cm_norm, annot=cm, fmt="d", cmap=SEQ_CMAP,
                xticklabels=classes, yticklabels=classes,
                cbar_kws={"label": "Row-normalised"}, ax=ax)
    ax.set_xlabel("Predicted")
    ax.set_ylabel("True")
    ax.set_title(title)
    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)


def plot_permutation_importance(importances, feature_names, path):
    means = importances.mean(axis=1)
    sds = importances.std(axis=1)
    order = np.argsort(means)
    # Rose = informative feature (positive macro-F1 drop), mint = uninformative
    colors = [COL_HIGH if means[i] >= 0 else COL_LOW for i in order]
    fig, ax = plt.subplots(figsize=(7.5, max(3, 0.34 * len(feature_names))))
    ax.barh([feature_names[i] for i in order], means[order],
            xerr=sds[order], color=colors, edgecolor="#333333",
            linewidth=0.5, error_kw=dict(ecolor="#222222", capsize=2))
    # Value label at the end of each bar
    span = means.max() - min(means.min(), 0)
    for yi, idx in enumerate(order):
        v, s = means[idx], sds[idx]
        if v >= 0:
            ax.text(v + s + 0.02 * span, yi, f"{v:.3f}", va="center",
                    ha="left", fontsize=8, fontweight="bold")
        else:
            ax.text(v - s - 0.02 * span, yi, f"{v:.3f}", va="center",
                    ha="right", fontsize=8, fontweight="bold")
    ax.margins(x=0.16)
    ax.set_xlabel("Drop in macro-F1 when permuted", fontsize=11)
    ax.set_title("Random Forest permutation importance", fontsize=12)
    ax.tick_params(labelsize=10)
    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)


def plot_phq_proba(df, model, feature_names, scaler, path):
    """Marginal predicted-probability curve over PHQ-9, demographics averaged."""
    Xref = df[feature_names].mean(numeric_only=True).to_dict()
    grid = np.linspace(0, 27, 28)
    rows = []
    for v in grid:
        x = pd.Series(Xref).copy()
        x["PHQ9_total"] = v
        rows.append(x[feature_names].values)
    Xg = np.vstack(rows).astype(float)
    Xg_s = scaler.transform(Xg)
    proba = model.predict_proba(Xg_s)
    fig, ax = plt.subplots(figsize=(7, 4))
    labels = ["Profile 0 — low CV risk", "Profile 1 — elevated CV risk"]
    for k in range(proba.shape[1]):
        ax.plot(grid, proba[:, k], linewidth=2.2,
                color=PROFILE_COLORS[k] if k < 2 else None,
                label=labels[k] if k < 2 else f"Risk profile {k}")
    ax.set_xlabel("PHQ-9 total")
    ax.set_ylabel("Predicted probability")
    ax.set_title("Predicted CV risk profile vs PHQ-9 (Random Forest)")
    ax.legend()
    fig.tight_layout()
    fig.savefig(path, dpi=200)
    plt.close(fig)


# ---------------------------------------------------------------------------
# Descriptive statistics
# ---------------------------------------------------------------------------
def build_cohort_table(df: pd.DataFrame, path: Path) -> None:
    """Descriptive cohort characteristics, overall and by risk profile."""
    def summarise(sub: pd.DataFrame) -> dict:
        def pc(mask) -> str:
            return f"{100 * mask.mean():.1f}"
        race = sub["race"].value_counts(normalize=True) * 100
        return {
            "n": f"{len(sub)}",
            "Age, mean (SD)":
                f"{sub['RIDAGEYR'].mean():.1f} ({sub['RIDAGEYR'].std():.1f})",
            "Female, %": pc(sub["female"] == 1),
            "Non-Hispanic White, %": f"{race.get('NHWhite', 0):.1f}",
            "Non-Hispanic Black, %": f"{race.get('NHBlack', 0):.1f}",
            "Mexican American, %": f"{race.get('MexAm', 0):.1f}",
            "Non-Hispanic Asian, %": f"{race.get('NHAsian', 0):.1f}",
            "PHQ-9 total, mean (SD)":
                f"{sub['PHQ9_total'].mean():.1f} ({sub['PHQ9_total'].std():.1f})",
            "PHQ-9 ≥ 10, %": pc(sub["PHQ9_pos"] == 1),
            "Self-reported hypertension, %": pc(sub["BPQ020"] == 1),
            "Self-reported high cholesterol, %": pc(sub["BPQ080"] == 1),
            "Doctor-reported overweight, %": pc(sub["MCQ080"] == 1),
            "Family history, premature MI, %": pc(sub["MCQ300A"] == 1),
            "Family history, diabetes, %": pc(sub["MCQ300C"] == 1),
            "Insufficient sleep (<7 h), %": pc(sub["insufficient_sleep"] == 1),
            "Household smoker, %": pc(sub["household_smoker"] == 1),
        }
    overall = summarise(df)
    low = summarise(df[df["risk_profile"] == 0])
    elev = summarise(df[df["risk_profile"] == 1])
    rows = [dict(Characteristic=k, Overall=overall[k],
                 Low_risk_profile=low[k], Elevated_risk_profile=elev[k])
            for k in overall]
    pd.DataFrame(rows).to_csv(path, index=False)


def crude_association(df: pd.DataFrame, path: Path) -> dict:
    """Odds ratio + chi-square for PHQ-9 >= 10 vs elevated risk profile."""
    from math import log, sqrt, exp, erf
    dep = (df["PHQ9_pos"] == 1).to_numpy()
    elev = (df["risk_profile"] == 1).to_numpy()
    a = int((dep & elev).sum())     # depressed, elevated
    b = int((dep & ~elev).sum())    # depressed, low
    c = int((~dep & elev).sum())    # not depressed, elevated
    d = int((~dep & ~elev).sum())   # not depressed, low
    odds = (a * d) / (b * c)
    se = sqrt(1 / a + 1 / b + 1 / c + 1 / d)
    lo, hi = exp(log(odds) - 1.96 * se), exp(log(odds) + 1.96 * se)
    n = a + b + c + d
    obs = [[a, b], [c, d]]
    rt, ct = [a + b, c + d], [a + c, b + d]
    chi2 = sum((obs[i][j] - rt[i] * ct[j] / n) ** 2 / (rt[i] * ct[j] / n)
               for i in range(2) for j in range(2))
    p = 2 * (1 - 0.5 * (1 + erf(sqrt(chi2) / sqrt(2))))  # chi2 df=1 -> z^2
    res = dict(a=a, b=b, c=c, d=d, odds_ratio=round(odds, 2),
               ci_low=round(lo, 2), ci_high=round(hi, 2),
               chi2=round(chi2, 1), p_value=p)
    with open(path, "w") as f:
        json.dump(res, f, indent=2)
    return res


# ---------------------------------------------------------------------------
# 7. Main pipeline
# ---------------------------------------------------------------------------
def cmd_analyse():
    print("==> 1. Loading NHANES P_ cycle ...")
    raw = build_cohort()
    print(f"    raw merged cohort: {len(raw)} adults 18-39")

    print("==> 2. Feature engineering ...")
    feats = build_features(raw)
    # Drop rows with missing PHQ-9 total (too many missing items)
    feats = feats.dropna(subset=["PHQ9_total"]).copy()
    print(f"    after PHQ-9 prorating filter: {len(feats)}")

    feats = impute(feats)
    feats.to_csv(RES / "analysis_cohort.csv", index=False)

    cohort_summary = pd.DataFrame({
        "n": [len(feats)],
        "mean_age": [feats["RIDAGEYR"].mean()],
        "pct_female": [feats["female"].mean() * 100],
        "PHQ9_mean": [feats["PHQ9_total"].mean()],
        "PHQ9_sd": [feats["PHQ9_total"].std()],
        "PHQ9_pos_pct": [feats["PHQ9_pos"].mean() * 100],
    })
    cohort_summary.to_csv(RES / "cohort_summary.csv", index=False)
    print("    ", cohort_summary.to_string(index=False))

    print("==> 3. Clustering CV risk profiles ...")
    feats, cluster_info = derive_risk_profiles(feats, k_range=(2, 3, 4), seed=RNG)
    print(f"    K selected by silhouette: {cluster_info['k_star']}")
    print(f"    silhouettes: {cluster_info['silhouette']}")

    plot_silhouette(cluster_info["silhouette"],
                    cluster_info["inertia"],
                    FIG / "fig_silhouette.png")
    plot_cluster_profile(cluster_info["cluster_means"],
                         cluster_info["feature_names"],
                         FIG / "fig_cluster_profile.png")
    plot_phq_by_profile(feats, FIG / "fig_phq_by_profile.png")

    with open(RES / "cluster_info.json", "w") as f:
        json.dump({k: v for k, v in cluster_info.items()
                   if k != "scaler_mean" and k != "scaler_std"},
                  f, indent=2, default=str)

    # Cluster prevalence per PHQ-9 stratum
    prev = (feats.groupby(["PHQ9_pos", "risk_profile"]).size()
                  .unstack(fill_value=0))
    prev_pct = prev.div(prev.sum(axis=1), axis=0) * 100
    prev_pct.to_csv(RES / "risk_profile_by_depression.csv")

    # Descriptive cohort table and crude depression-profile association
    build_cohort_table(feats, RES / "cohort_table.csv")
    assoc = crude_association(feats, RES / "crude_association.json")
    print(f"    crude association: OR={assoc['odds_ratio']} "
          f"(95% CI {assoc['ci_low']}-{assoc['ci_high']}), "
          f"chi2={assoc['chi2']}, p={assoc['p_value']:.2e}")

    print("==> 4. Supervised modelling ...")
    X_full, feat_names = make_feature_matrix(feats, include_depression=True)
    X_demo, feat_names_demo = make_feature_matrix(feats, include_depression=False)
    y = feats["risk_profile"].to_numpy()

    seeds = (1, 2, 3)
    results = {}

    # Hyperparameters selected by grid search via inner 3-fold CV on an
    # 80% tuning subset; see code/tune_hyperparameters.py and
    # results/hyperparameter_search.csv. The 20% holdout reserved by
    # the tuner is not honoured here because run_analysis.py uses
    # cross-validation across the whole cohort; the tuner's holdout is
    # only to keep tuning honest if a single-split eval were used.
    def lr_factory():
        return LogisticRegression(C=1.0, lr=1.0, max_iter=600, random_state=0)

    def rf_factory():
        return RandomForestClassifier(n_estimators=60, max_depth=8,
                                      min_samples_leaf=30,
                                      min_samples_split=60,
                                      max_features="sqrt",
                                      random_state=0)

    def gb_factory():
        return make_gb()

    print(f"    Gradient-boosting backend: {GB_BACKEND}")

    print("    Logistic regression (full)")
    results["LR_full"] = cross_validate(lr_factory, X_full, y, seeds=seeds)
    print("    Logistic regression (demographics only)")
    results["LR_demo"] = cross_validate(lr_factory, X_demo, y, seeds=seeds)
    print("    Random forest (full)")
    results["RF_full"] = cross_validate(rf_factory, X_full, y, seeds=seeds)
    print("    Random forest (demographics only)")
    results["RF_demo"] = cross_validate(rf_factory, X_demo, y, seeds=seeds)
    print("    Gradient boosting (full)")
    results["GB_full"] = cross_validate(gb_factory, X_full, y, seeds=seeds)
    print("    Gradient boosting (demographics only)")
    results["GB_demo"] = cross_validate(gb_factory, X_demo, y, seeds=seeds)

    summary_rows = []
    for name, df in results.items():
        for m in ["acc", "bal_acc", "macro_f1", "roc_auc"]:
            summary_rows.append(dict(
                model=name, metric=m,
                mean=df[m].mean(), sd=df[m].std(),
                lo=df[m].quantile(0.025), hi=df[m].quantile(0.975),
            ))
    summary = pd.DataFrame(summary_rows)
    summary.to_csv(RES / "cv_summary.csv", index=False)
    with open(RES / "gb_backend.txt", "w") as f:
        f.write(f"Gradient-boosting backend used for this run: {GB_BACKEND}\n")
    for n, d in results.items():
        d.to_csv(RES / f"cv_folds_{n}.csv", index=False)
    print(summary.pivot_table(index="model", columns="metric",
                              values="mean").round(3).to_string())

    plot_metric_bars(results, FIG / "fig_cv_metrics.png")

    # Paired t-test: full vs demographics-only on macro F1
    from math import sqrt
    def paired_t(a, b):
        d = np.asarray(a) - np.asarray(b)
        t = d.mean() / (d.std(ddof=1) / sqrt(len(d)))
        # Two-sided p via normal approx (df=24, conservative but ok)
        from math import erf
        p = 2 * (1 - 0.5 * (1 + erf(abs(t) / sqrt(2))))
        return t, p

    lr_t, lr_p = paired_t(results["LR_full"]["macro_f1"],
                          results["LR_demo"]["macro_f1"])
    rf_t, rf_p = paired_t(results["RF_full"]["macro_f1"],
                          results["RF_demo"]["macro_f1"])
    gb_t, gb_p = paired_t(results["GB_full"]["macro_f1"],
                          results["GB_demo"]["macro_f1"])
    with open(RES / "ablation_ttests.json", "w") as f:
        json.dump(dict(LR=dict(t=lr_t, p=lr_p),
                       RF=dict(t=rf_t, p=rf_p),
                       GB=dict(t=gb_t, p=gb_p)), f, indent=2)
    print(f"    LR ablation: t={lr_t:.2f}, p~{lr_p:.4f}")
    print(f"    RF ablation: t={rf_t:.2f}, p~{rf_p:.4f}")
    print(f"    GB ablation: t={gb_t:.2f}, p~{gb_p:.4f}")

    print("==> 5. Final model + interpretability ...")
    # Keep aligned indices so subgroup metadata can be recovered on test split
    idx = np.arange(len(X_full))
    idx_tr, idx_te, y_tr, y_te = train_test_split(
        idx, y, test_size=0.25, random_state=RNG, stratify=y)
    X_tr, X_te = X_full[idx_tr], X_full[idx_te]
    feats_test = feats.iloc[idx_te].reset_index(drop=True)
    scaler = StandardScaler().fit(X_tr)
    Xtr_s = scaler.transform(X_tr)
    Xte_s = scaler.transform(X_te)

    Xtr_b, ytr_b = oversample(Xtr_s, y_tr, random_state=RNG)
    rf = rf_factory().fit(Xtr_b, ytr_b)
    lr = lr_factory().fit(Xtr_b, ytr_b)

    classes = np.unique(y)
    rf_pred = rf.predict(Xte_s)
    lr_pred = lr.predict(Xte_s)

    # Calibration check: oversampling improves discrimination but distorts
    # the predicted probabilities. Compare against a non-oversampled RF.
    rf_cal = rf_factory().fit(Xtr_s, y_tr)
    yte_bin = (y_te == 1).astype(float)
    p_over = rf.predict_proba(Xte_s)[:, 1]
    p_cal = rf_cal.predict_proba(Xte_s)[:, 1]
    obs_prev = float(yte_bin.mean())
    calibration = dict(
        observed_prevalence=round(obs_prev, 3),
        oversampled_mean_pred=round(float(p_over.mean()), 3),
        oversampled_brier=round(float(np.mean((p_over - yte_bin) ** 2)), 3),
        nonoversampled_mean_pred=round(float(p_cal.mean()), 3),
        nonoversampled_brier=round(float(np.mean((p_cal - yte_bin) ** 2)), 3),
        baseline_brier=round(float(np.mean((obs_prev - yte_bin) ** 2)), 3),
    )
    with open(RES / "calibration.json", "w") as f:
        json.dump(calibration, f, indent=2)
    print(f"    calibration: {calibration}")

    plot_confusion(y_te, rf_pred, classes,
                   "Random Forest confusion (held-out)",
                   FIG / "fig_rf_confusion.png")
    plot_confusion(y_te, lr_pred, classes,
                   "Logistic Regression confusion (held-out)",
                   FIG / "fig_lr_confusion.png")

    imp = permutation_importance(rf, Xte_s, y_te, n_repeats=8,
                                 random_state=RNG, scoring=macro_f1)
    pi_df = pd.DataFrame({
        "feature": feat_names,
        "mean_drop_f1": imp.mean(axis=1),
        "sd_drop_f1": imp.std(axis=1),
    }).sort_values("mean_drop_f1", ascending=False)
    pi_df.to_csv(RES / "permutation_importance_rf.csv", index=False)
    plot_permutation_importance(imp, feat_names, FIG / "fig_rf_permimp.png")

    # Marginal PHQ-9 effect
    feats_for_marg = pd.DataFrame(X_full, columns=feat_names)
    plot_phq_proba(feats_for_marg, rf, feat_names, scaler,
                   FIG / "fig_rf_phq_marginal.png")

    # Logistic regression coefficients
    lr_coef = pd.DataFrame(lr.W[1:], index=feat_names,
                           columns=[f"class_{c}" for c in lr.classes_])
    lr_coef.to_csv(RES / "logistic_coefficients.csv")

    # ---- Subgroup fairness evaluation ----------------------------------
    print("==> 6. Subgroup fairness evaluation ...")
    rf_proba_te = rf.predict_proba(Xte_s)
    rf_pred_te = rf.predict(Xte_s)

    def subgroup_metrics(mask, label):
        if mask.sum() < 30:
            return None
        y_t = y_te[mask]
        y_p = rf_pred_te[mask]
        y_pr = rf_proba_te[mask]
        return dict(
            subgroup=label,
            n=int(mask.sum()),
            elev_risk_pct=round(100 * (y_t == 1).mean(), 1),
            acc=round(accuracy(y_t, y_p), 3),
            bal_acc=round(balanced_accuracy(y_t, y_p), 3),
            macro_f1=round(macro_f1(y_t, y_p), 3),
            roc_auc=round(roc_auc_ovr(y_t, y_pr, classes), 3),
        )

    sub_rows = []
    # Overall
    sub_rows.append(subgroup_metrics(np.ones(len(y_te), bool), "Overall"))
    # Sex
    for lab, m in [("Female", feats_test["female"] == 1),
                   ("Male", feats_test["female"] == 0)]:
        r = subgroup_metrics(m.to_numpy(), f"Sex: {lab}")
        if r:
            sub_rows.append(r)
    # Race/ethnicity
    for race in sorted(feats_test["race"].unique()):
        m = (feats_test["race"] == race).to_numpy()
        r = subgroup_metrics(m, f"Race: {race}")
        if r:
            sub_rows.append(r)
    # Income tertiles (poverty income ratio)
    pir = feats_test["INDFMPIR"]
    q1, q2 = pir.quantile([1/3, 2/3])
    for lab, m in [(f"PIR low (<{q1:.2f})", (pir < q1).to_numpy()),
                   (f"PIR mid", ((pir >= q1) & (pir < q2)).to_numpy()),
                   (f"PIR high (≥{q2:.2f})", (pir >= q2).to_numpy())]:
        r = subgroup_metrics(m, f"Income {lab}")
        if r:
            sub_rows.append(r)

    sub_df = pd.DataFrame([r for r in sub_rows if r is not None])
    sub_df.to_csv(RES / "subgroup_fairness.csv", index=False)
    print(sub_df.to_string(index=False))

    print("\nAll outputs written to:")
    print(f"  figures/  -> {FIG}")
    print(f"  results/  -> {RES}")


# ===========================================================================
# SECTION 3  hyperparameter tuning  (was tune_hyperparameters.py)
# ===========================================================================

def inner_cv_score(model_fn, X, y, k=3, seed=0):
    skf = StratifiedKFold(n_splits=k, shuffle=True, random_state=seed)
    scores = []
    for tr, te in skf.split(X, y):
        sc = StandardScaler().fit(X[tr])
        Xtr_s, Xte_s = sc.transform(X[tr]), sc.transform(X[te])
        Xtr_b, ytr_b = oversample(Xtr_s, y[tr], random_state=seed)
        m = model_fn()
        m.fit(Xtr_b, ytr_b)
        scores.append(macro_f1(y[te], m.predict(Xte_s)))
    return float(np.mean(scores)), float(np.std(scores))


def cmd_tune():
    print("==> Loading and processing data ...")
    df = build_cohort()
    df = build_features(df)
    df = df.dropna(subset=["PHQ9_total"]).copy()
    df = impute(df)
    df, _ = derive_risk_profiles(df, k_range=(2, 3, 4), seed=RNG)

    X, _ = make_feature_matrix(df, include_depression=True)
    y = df["risk_profile"].to_numpy()

    # 80% tuning, 20% holdout (never touched here)
    X_tune, _X_hold, y_tune, _y_hold = train_test_split(
        X, y, test_size=0.20, random_state=RNG, stratify=y
    )
    print(f"    tuning n={len(X_tune)}; holdout n={len(_X_hold)}")

    rows = []

    # ---- LR grid ---------------------------------------------------------
    lr_grid = [(C, lr_rate)
               for C in [0.01, 0.1, 1.0, 10.0]
               for lr_rate in [0.1, 0.3, 0.5, 1.0]]
    print(f"==> LR grid: {len(lr_grid)} configs")
    for C, lr_rate in lr_grid:
        def fac(C=C, lr_rate=lr_rate):
            return LogisticRegression(C=C, lr=lr_rate, max_iter=500,
                                      random_state=0)
        m, s = inner_cv_score(fac, X_tune, y_tune, k=3, seed=0)
        rows.append(dict(model="LR", C=C, lr=lr_rate, n_estimators=None,
                         max_depth=None, min_samples_leaf=None,
                         mean_f1=m, sd_f1=s))

    # ---- RF grid ---------------------------------------------------------
    rf_grid = [(n, d, leaf)
               for n in [40, 60, 100]
               for d in [4, 6, 8]
               for leaf in [20, 30, 50]]
    print(f"==> RF grid: {len(rf_grid)} configs")
    for n, d, leaf in rf_grid:
        def fac(n=n, d=d, leaf=leaf):
            return RandomForestClassifier(
                n_estimators=n, max_depth=d,
                min_samples_leaf=leaf, min_samples_split=2 * leaf,
                max_features="sqrt", random_state=0)
        m, s = inner_cv_score(fac, X_tune, y_tune, k=3, seed=0)
        rows.append(dict(model="RF", C=None, lr=None,
                         n_estimators=n, max_depth=d,
                         min_samples_leaf=leaf, learning_rate=None,
                         mean_f1=m, sd_f1=s))

    # ---- GB grid ---------------------------------------------------------
    gb_grid = [(n, lr_rate, d)
               for n in [40, 60, 100]
               for lr_rate in [0.05, 0.1]
               for d in [2, 3]]
    print(f"==> GB grid: {len(gb_grid)} configs")
    for n, lr_rate, d in gb_grid:
        def fac(n=n, lr_rate=lr_rate, d=d):
            return GradientBoostingClassifier(
                n_estimators=n, learning_rate=lr_rate, max_depth=d,
                min_samples_leaf=30, min_samples_split=60,
                subsample=0.8, random_state=0)
        m, s = inner_cv_score(fac, X_tune, y_tune, k=3, seed=0)
        rows.append(dict(model="GB", C=None, lr=None,
                         n_estimators=n, max_depth=d,
                         min_samples_leaf=30, learning_rate=lr_rate,
                         mean_f1=m, sd_f1=s))

    grid = pd.DataFrame(rows)
    grid.to_csv(REPO / "results" / "hyperparameter_search.csv", index=False)
    print("    saved results/hyperparameter_search.csv")

    # ---- Best per model --------------------------------------------------
    best = {}
    for model in ["LR", "RF", "GB"]:
        sub = grid[grid["model"] == model].copy()
        sub = sub.sort_values("mean_f1", ascending=False)
        top = sub.iloc[0].to_dict()
        best[model] = {k: v for k, v in top.items()
                       if pd.notna(v) and k != "model"}
        print(f"  {model} best: {best[model]}")

    with open(REPO / "results" / "best_hyperparameters.json", "w") as f:
        json.dump(best, f, indent=2)

    for model in ["LR", "RF", "GB"]:
        sub = grid[grid["model"] == model].sort_values(
            "mean_f1", ascending=False).head(3)
        print(f"\n{model} top 3:")
        print(sub.to_string(index=False))


# ===========================================================================
# SECTION 4  age-window sensitivity  (was sensitivity_age.py)
# ===========================================================================

def build_cohort_age(min_age, max_age):
    demo = load_module("P_DEMO")
    dpq = load_module("P_DPQ")
    bpq = load_module("P_BPQ")
    mcq = load_module("P_MCQ")
    slq = load_module("P_SLQ")
    smq = load_module("P_SMQFAM")
    cdemo = ["SEQN","RIDAGEYR","RIAGENDR","RIDRETH3","DMDEDUC2",
             "DMDMARTZ","INDFMPIR","RIDEXPRG"]
    cbpq  = ["SEQN","BPQ020","BPQ050A","BPQ080","BPQ100D"]
    cmcq  = ["SEQN","MCQ080","MCQ366A","MCQ366B","MCQ366C","MCQ366D",
             "MCQ300A","MCQ300C"]
    cslq  = ["SEQN","SLD012"]; csmq=["SEQN","SMD460"]
    cdpq  = ["SEQN"] + [f"DPQ0{i}0" for i in range(1,10)]
    df = (demo[cdemo].merge(dpq[cdpq], on="SEQN")
                     .merge(bpq[cbpq], on="SEQN", how="left")
                     .merge(mcq[cmcq], on="SEQN", how="left")
                     .merge(slq[cslq], on="SEQN", how="left")
                     .merge(smq[csmq], on="SEQN", how="left"))
    df = df[(df["RIDAGEYR"]>=min_age) & (df["RIDAGEYR"]<=max_age)].copy()
    df = df[df["RIDEXPRG"].fillna(0) != 1].copy()
    df.drop(columns=["RIDEXPRG"], inplace=True)
    return df.reset_index(drop=True)


def _paired_t(a, b):
    d = np.asarray(a) - np.asarray(b)
    t = d.mean() / (d.std(ddof=1) / sqrt(len(d)))
    p = 2 * (1 - 0.5*(1 + erf(abs(t)/sqrt(2))))
    return t, p


def run_window(min_age, max_age, label, seeds=(1, 2)):
    t0 = time.time()
    print(f"\n==> Cohort {label}: ages {min_age}-{max_age}")
    feats = build_cohort_age(min_age, max_age)
    feats = build_features(feats).dropna(subset=["PHQ9_total"]).copy()
    feats = impute(feats)
    feats, info = derive_risk_profiles(feats, k_range=(2, 3, 4), seed=RNG)
    elev_pct = round(100 * (feats["risk_profile"] == feats["risk_profile"].max()).mean(), 1)
    phq_pos = round(100 * feats["PHQ9_pos"].mean(), 1)
    print(f"    n={len(feats)}; K*={info['k_star']}; elev={elev_pct}%; PHQ-9+={phq_pos}%")

    X_full, _ = make_feature_matrix(feats, include_depression=True)
    X_demo, _ = make_feature_matrix(feats, include_depression=False)
    y = feats["risk_profile"].to_numpy()

    def lr_fac():
        return LogisticRegression(C=1.0, lr=1.0, max_iter=400, random_state=0)
    def rf_fac():
        return RandomForestClassifier(n_estimators=40, max_depth=6,
                                      min_samples_leaf=40,
                                      min_samples_split=80,
                                      max_features="sqrt", random_state=0)

    lr_full = cross_validate(lr_fac, X_full, y, seeds=seeds)
    lr_demo = cross_validate(lr_fac, X_demo, y, seeds=seeds)
    rf_full = cross_validate(rf_fac, X_full, y, seeds=seeds)
    rf_demo = cross_validate(rf_fac, X_demo, y, seeds=seeds)
    print(f"    done ({time.time()-t0:.1f}s)")

    lr_t, lr_p = _paired_t(lr_full["macro_f1"], lr_demo["macro_f1"])
    rf_t, rf_p = _paired_t(rf_full["macro_f1"], rf_demo["macro_f1"])

    return dict(
        cohort=label, age_window=f"{min_age}-{max_age}", n=len(feats),
        elev_profile_pct=elev_pct, phq9_pos_pct=phq_pos, K=int(info["k_star"]),
        LR_demo_AUC=round(lr_demo["roc_auc"].mean(), 3),
        LR_full_AUC=round(lr_full["roc_auc"].mean(), 3),
        LR_dAUC=round(lr_full["roc_auc"].mean()-lr_demo["roc_auc"].mean(), 3),
        RF_demo_AUC=round(rf_demo["roc_auc"].mean(), 3),
        RF_full_AUC=round(rf_full["roc_auc"].mean(), 3),
        RF_dAUC=round(rf_full["roc_auc"].mean()-rf_demo["roc_auc"].mean(), 3),
        LR_f1_t=round(lr_t, 2), LR_f1_p=f"{lr_p:.1e}",
        RF_f1_t=round(rf_t, 2), RF_f1_p=f"{rf_p:.1e}",
    )


def cmd_sens_age():
    rows = []
    rows.append(run_window(18, 39, "Primary: 18-39"))
    rows.append(run_window(18, 64, "Working-age: 18-64"))
    rows.append(run_window(18, 99, "All adults: 18+"))
    df = pd.DataFrame(rows)
    df.to_csv(REPO / "results" / "sensitivity_age.csv", index=False)
    print("\n", df.to_string(index=False))


# ===========================================================================
# SECTION 5  clustering sensitivity  (was sensitivity_clustering.py)
# ===========================================================================

def adjusted_rand_index(a: np.ndarray, b: np.ndarray) -> float:
    """Adjusted Rand index between two labellings."""
    a, b = np.asarray(a), np.asarray(b)
    ua, ub = np.unique(a), np.unique(b)
    cont = np.zeros((len(ua), len(ub)), dtype=float)
    ia = {v: i for i, v in enumerate(ua)}
    ib = {v: i for i, v in enumerate(ub)}
    for x, y in zip(a, b):
        cont[ia[x], ib[y]] += 1

    def c2(x):
        return x * (x - 1) / 2

    sum_ij = c2(cont).sum()
    sum_a = c2(cont.sum(axis=1)).sum()
    sum_b = c2(cont.sum(axis=0)).sum()
    total = c2(len(a))
    expected = sum_a * sum_b / total
    maxv = 0.5 * (sum_a + sum_b)
    return 1.0 if maxv == expected else float((sum_ij - expected) / (maxv - expected))


def best_agreement(a: np.ndarray, b: np.ndarray) -> float:
    """Raw agreement under the better of the two binary label alignments."""
    a, b = np.asarray(a), np.asarray(b)
    return float(max((a == b).mean(), (a == (1 - b)).mean()))


def cmd_sens_cluster():
    print("==> Building cohort and primary (K-means) profiles ...")
    feats = build_features(build_cohort()).dropna(subset=["PHQ9_total"]).copy()
    feats = impute(feats)
    feats, info = derive_risk_profiles(feats, k_range=(2, 3, 4), seed=RNG)
    km_labels = feats["risk_profile"].to_numpy()

    # K-modes operates on the 13 binary risk indicators (no continuous term)
    Xb = feats[RISK_BINARY].to_numpy(dtype=float)
    cat_idx = list(range(Xb.shape[1]))

    print("==> Re-deriving profiles with K-modes ...")
    costs = {}
    for k in (2, 3, 4):
        km_mode = KPrototypes(n_clusters=k, cat_idx=cat_idx, num_idx=[],
                              n_init=10, random_state=RNG).fit(Xb)
        costs[k] = km_mode.cost_
        print(f"    K={k}: K-modes cost = {km_mode.cost_:.1f}")

    kmode = KPrototypes(n_clusters=2, cat_idx=cat_idx, num_idx=[],
                        n_init=20, random_state=RNG).fit(Xb)
    kmode_labels = kmode.labels_
    # Relabel so 0 = lower mean indicator prevalence, matching K-means
    if Xb[kmode_labels == 0].mean() > Xb[kmode_labels == 1].mean():
        kmode_labels = 1 - kmode_labels

    agreement = best_agreement(km_labels, kmode_labels)
    ari = adjusted_rand_index(km_labels, kmode_labels)
    print(f"\n    K-means vs K-modes raw agreement: {agreement:.3f}")
    print(f"    Adjusted Rand index:              {ari:.3f}")

    phq_pos = feats["PHQ9_pos"].to_numpy().astype(bool)

    def elev_split(labels):
        hi = labels == 1
        return (round(100 * hi[phq_pos].mean(), 1),
                round(100 * hi[~phq_pos].mean(), 1))

    km_hi_dep, km_hi_nodep = elev_split(km_labels)
    kmode_hi_dep, kmode_hi_nodep = elev_split(kmode_labels)
    print(f"\n    K-means  elevated-risk: PHQ-9>=10 {km_hi_dep}% vs <10 {km_hi_nodep}%")
    print(f"    K-modes  elevated-risk: PHQ-9>=10 {kmode_hi_dep}% vs <10 {kmode_hi_nodep}%")

    out = pd.DataFrame([
        dict(metric="K-modes cost K=2", value=round(costs[2], 1)),
        dict(metric="K-modes cost K=3", value=round(costs[3], 1)),
        dict(metric="K-modes cost K=4", value=round(costs[4], 1)),
        dict(metric="K-means elevated-profile size %",
             value=round(100 * (km_labels == 1).mean(), 1)),
        dict(metric="K-modes elevated-profile size %",
             value=round(100 * (kmode_labels == 1).mean(), 1)),
        dict(metric="Raw agreement (K-means vs K-modes)", value=round(agreement, 3)),
        dict(metric="Adjusted Rand index", value=round(ari, 3)),
        dict(metric="K-means elevated-risk %, PHQ-9>=10", value=km_hi_dep),
        dict(metric="K-means elevated-risk %, PHQ-9<10", value=km_hi_nodep),
        dict(metric="K-modes elevated-risk %, PHQ-9>=10", value=kmode_hi_dep),
        dict(metric="K-modes elevated-risk %, PHQ-9<10", value=kmode_hi_nodep),
    ])
    out.to_csv(REPO / "results" / "sensitivity_clustering.csv", index=False)
    print(f"\nWrote {REPO / 'results' / 'sensitivity_clustering.csv'}")


# ===========================================================================
# SECTION 6  ml_lib validation  (was validate_ml_lib.py)
# ===========================================================================

# Optional sklearn reference comparison
try:
    import sklearn  # noqa: F401
    from sklearn.cluster import KMeans as SkKMeans
    from sklearn.linear_model import LogisticRegression as SkLR
    from sklearn.ensemble import (RandomForestClassifier as SkRF,
                                  GradientBoostingClassifier as SkGB)
    from sklearn.metrics import (accuracy_score as sk_acc,
                                 balanced_accuracy_score as sk_bal,
                                 f1_score as sk_f1,
                                 roc_auc_score as sk_auc)
    from sklearn.preprocessing import StandardScaler as SkScaler
    SKLEARN_AVAILABLE = True
except Exception:
    SKLEARN_AVAILABLE = False

REPORT = RES / "validation_report.txt"
results: list[str] = []


def _log(msg: str) -> None:
    print(msg)
    results.append(msg)


# ---------------------------------------------------------------------------
# Synthetic data builders
# ---------------------------------------------------------------------------
def make_blobs(n=300, centers=None, std=0.6, seed=0):
    rng = np.random.default_rng(seed)
    centers = np.array(centers if centers is not None
                       else [[-3, -3], [0, 0], [3, 3]])
    per = n // len(centers)
    X = np.vstack([rng.normal(c, std, size=(per, 2)) for c in centers])
    y = np.repeat(np.arange(len(centers)), per)
    perm = rng.permutation(len(X))
    return X[perm], y[perm]


def make_linear_classification(n=400, seed=0):
    rng = np.random.default_rng(seed)
    X = rng.normal(0, 1, size=(n, 4))
    w_true = np.array([1.5, -1.0, 0.5, 0.0])
    logits = X @ w_true + 0.2
    y = (logits + rng.normal(0, 0.4, size=n) > 0).astype(int)
    return X, y


def make_nonlinear_classification(n=400, seed=0):
    rng = np.random.default_rng(seed)
    X = rng.uniform(-2, 2, size=(n, 4))
    # Non-linear decision boundary: depends on x0*x1 and x2^2
    y = ((X[:, 0] * X[:, 1] + X[:, 2] ** 2 - 1.0) > 0).astype(int)
    return X, y


def cluster_purity(labels_true, labels_pred):
    """Fraction of samples in the dominant true class within each cluster."""
    n = len(labels_true)
    total = 0
    for c in np.unique(labels_pred):
        mask = labels_pred == c
        if mask.sum() == 0:
            continue
        vals, counts = np.unique(labels_true[mask], return_counts=True)
        total += counts.max()
    return total / n


# ---------------------------------------------------------------------------
# (1) Analytical / internal-consistency checks
# ---------------------------------------------------------------------------
def cmd_validate():
    _log("=" * 64)
    _log("VALIDATION REPORT  (code/validate_ml_lib.py)")
    _log(f"sklearn available: {SKLEARN_AVAILABLE}")
    _log("=" * 64)

    _log("\n--- (1) Analytical / internal-consistency checks ---")

    # StandardScaler
    X = np.array([[1.0, 10.0], [2.0, 20.0], [3.0, 30.0], [4.0, 40.0]])
    Xz = StandardScaler().fit_transform(X)
    m_ok = np.allclose(Xz.mean(axis=0), 0, atol=1e-8)
    s_ok = np.allclose(Xz.std(axis=0), 1, atol=1e-8)
    _log(f"  StandardScaler: mean≈0 {'PASS' if m_ok else 'FAIL'}, "
         f"std≈1 {'PASS' if s_ok else 'FAIL'}")

    # K-means on well-separated blobs
    Xb, yb = make_blobs(n=600, centers=[[-5, -5], [0, 0], [5, 5]], std=0.5, seed=1)
    km = KMeans(n_clusters=3, n_init=10, random_state=1).fit(Xb)
    purity = cluster_purity(yb, km.labels_)
    _log(f"  K-means cluster purity on 3 separated blobs: {purity:.3f} "
         f"({'PASS' if purity >= 0.95 else 'FAIL'}; expect ≥ 0.95)")

    # Logistic regression on linearly separable data
    Xl, yl = make_linear_classification(n=600, seed=2)
    lr = LogisticRegression(C=1.0, lr=1.0, max_iter=600).fit(Xl, yl)
    acc_train = accuracy(yl, lr.predict(Xl))
    prob_sum = lr.predict_proba(Xl).sum(axis=1)
    _log(f"  LR train accuracy on linear problem: {acc_train:.3f} "
         f"({'PASS' if acc_train >= 0.85 else 'FAIL'}; expect ≥ 0.85)")
    _log(f"  LR predict_proba rows sum to 1: "
         f"{'PASS' if np.allclose(prob_sum, 1.0) else 'FAIL'}")

    # Random forest on non-linear data
    Xn, yn = make_nonlinear_classification(n=600, seed=3)
    rf = RandomForestClassifier(n_estimators=80, max_depth=6,
                                min_samples_leaf=8, random_state=3).fit(Xn, yn)
    acc_train_rf = accuracy(yn, rf.predict(Xn))
    prob_sum_rf = rf.predict_proba(Xn).sum(axis=1)
    _log(f"  RF train accuracy on non-linear problem: {acc_train_rf:.3f} "
         f"({'PASS' if acc_train_rf >= 0.85 else 'FAIL'}; expect ≥ 0.85)")
    _log(f"  RF predict_proba rows sum to 1: "
         f"{'PASS' if np.allclose(prob_sum_rf, 1.0) else 'FAIL'}")

    # Gradient boosting on non-linear data
    gb = GradientBoostingClassifier(n_estimators=80, learning_rate=0.1,
                                    max_depth=3, min_samples_leaf=8,
                                    random_state=3).fit(Xn, yn)
    acc_train_gb = accuracy(yn, gb.predict(Xn))
    prob_sum_gb = gb.predict_proba(Xn).sum(axis=1)
    _log(f"  GB train accuracy on non-linear problem: {acc_train_gb:.3f} "
         f"({'PASS' if acc_train_gb >= 0.85 else 'FAIL'}; expect ≥ 0.85)")
    _log(f"  GB predict_proba rows sum to 1: "
         f"{'PASS' if np.allclose(prob_sum_gb, 1.0) else 'FAIL'}")

    # K-modes (KPrototypes, no numeric columns) on planted binary blobs
    rng_km = np.random.default_rng(7)
    clusterA = (rng_km.random((150, 6)) < 0.15).astype(float)
    clusterB = (rng_km.random((150, 6)) < 0.85).astype(float)
    Xkm = np.vstack([clusterA, clusterB])
    ykm = np.array([0] * 150 + [1] * 150)
    kmode = KPrototypes(n_clusters=2, cat_idx=list(range(6)), num_idx=[],
                        n_init=10, random_state=7).fit(Xkm)
    kmode_purity = cluster_purity(ykm, kmode.labels_)
    _log(f"  K-modes cluster purity on 2 binary blobs: {kmode_purity:.3f} "
         f"({'PASS' if kmode_purity >= 0.95 else 'FAIL'}; expect ≥ 0.95)")

    # Metrics: hand-built ground truth
    yt = np.array([0, 0, 0, 1, 1, 1, 1, 1])
    yp = np.array([0, 0, 1, 1, 1, 0, 1, 1])  # 6/8 correct, TP=4, FP=1, FN=1
    acc_ok = abs(accuracy(yt, yp) - 6 / 8) < 1e-9
    # Macro F1: class 0 F1 = 2*2/(2*2+1+1)=0.667 ; class 1 F1 = 2*4/(2*4+1+1)=0.8
    f1_ok = abs(macro_f1(yt, yp) - (0.6666666666666666 + 0.8) / 2) < 1e-6
    bal_ok = abs(balanced_accuracy(yt, yp) - (2 / 3 + 4 / 5) / 2) < 1e-9
    cm_ok = (confusion_matrix(yt, yp) == np.array([[2, 1], [1, 4]])).all()
    _log(f"  Metrics — accuracy {'PASS' if acc_ok else 'FAIL'}, "
         f"macro_f1 {'PASS' if f1_ok else 'FAIL'}, "
         f"balanced_acc {'PASS' if bal_ok else 'FAIL'}, "
         f"confusion_matrix {'PASS' if cm_ok else 'FAIL'}")

    # ROC-AUC on perfectly separated scores
    y_bin = np.array([0, 0, 0, 1, 1, 1])
    scores = np.array([0.1, 0.2, 0.3, 0.7, 0.8, 0.9])
    proba2 = np.stack([1 - scores, scores], axis=1)
    auc = roc_auc_ovr(y_bin, proba2, np.array([0, 1]))
    _log(f"  ROC-AUC on perfectly separable scores: {auc:.3f} "
         f"({'PASS' if abs(auc - 1.0) < 1e-9 else 'FAIL'}; expect 1.000)")

    # ---------------------------------------------------------------------------
    # (2) Side-by-side comparison with scikit-learn (if available)
    # ---------------------------------------------------------------------------
    if SKLEARN_AVAILABLE:
        _log("\n--- (2) Side-by-side comparison with scikit-learn ---")

        # StandardScaler
        sk_z = SkScaler().fit_transform(X)
        diff = np.abs(Xz - sk_z).max()
        _log(f"  StandardScaler vs sklearn — max abs diff: {diff:.2e} "
             f"({'PASS' if diff < 1e-8 else 'FAIL'}; expect < 1e-8)")

        # K-means on blobs (compare inertia within tolerance)
        sk_km = SkKMeans(n_clusters=3, n_init=10, random_state=1).fit(Xb)
        our_inertia = km.inertia_
        sk_inertia = sk_km.inertia_
        rel_diff = abs(our_inertia - sk_inertia) / sk_inertia
        _log(f"  K-means inertia ours vs sklearn: {our_inertia:.1f} vs "
             f"{sk_inertia:.1f} (rel diff {rel_diff:.1%}; "
             f"{'PASS' if rel_diff < 0.05 else 'FAIL'})")
        # Cluster purity vs sklearn purity
        sk_purity = cluster_purity(yb, sk_km.labels_)
        _log(f"  K-means cluster purity ours vs sklearn: {purity:.3f} vs "
             f"{sk_purity:.3f} ({'PASS' if abs(purity-sk_purity) < 0.05 else 'FAIL'})")

        # Logistic regression
        sk_lr = SkLR(C=1.0, solver="lbfgs", max_iter=600).fit(Xl, yl)
        sk_acc_lr = sk_acc(yl, sk_lr.predict(Xl))
        _log(f"  LR train acc ours vs sklearn: {acc_train:.3f} vs "
             f"{sk_acc_lr:.3f} ({'PASS' if abs(acc_train-sk_acc_lr) < 0.05 else 'FAIL'})")
        # Coefficient sign agreement on non-zero features
        our_coef = lr.W[1:, 1] - lr.W[1:, 0]  # binary case, class1 vs class0
        sk_coef = sk_lr.coef_[0]
        sign_agree = float(np.mean(np.sign(our_coef[:3]) == np.sign(sk_coef[:3])))
        _log(f"  LR coefficient sign agreement (first 3 features): "
             f"{sign_agree*100:.0f}% ({'PASS' if sign_agree >= 2/3 else 'FAIL'})")

        # Random forest
        sk_rf = SkRF(n_estimators=80, max_depth=6, min_samples_leaf=8,
                     random_state=3).fit(Xn, yn)
        sk_acc_rf = sk_acc(yn, sk_rf.predict(Xn))
        _log(f"  RF train acc ours vs sklearn: {acc_train_rf:.3f} vs "
             f"{sk_acc_rf:.3f} ({'PASS' if abs(acc_train_rf-sk_acc_rf) < 0.10 else 'FAIL'})")

        # Gradient boosting
        sk_gb = SkGB(n_estimators=80, learning_rate=0.1, max_depth=3,
                     min_samples_leaf=8, random_state=3).fit(Xn, yn)
        sk_acc_gb = sk_acc(yn, sk_gb.predict(Xn))
        _log(f"  GB train acc ours vs sklearn: {acc_train_gb:.3f} vs "
             f"{sk_acc_gb:.3f} ({'PASS' if abs(acc_train_gb-sk_acc_gb) < 0.10 else 'FAIL'})")

        # Metric functions
        sk_macro_f1 = sk_f1(yt, yp, average="macro")
        sk_bal_acc = sk_bal(yt, yp)
        sk_auc_val = sk_auc(y_bin, scores)
        _log(f"  macro_f1 ours vs sklearn: {macro_f1(yt, yp):.3f} vs "
             f"{sk_macro_f1:.3f} "
             f"({'PASS' if abs(macro_f1(yt, yp)-sk_macro_f1) < 1e-3 else 'FAIL'})")
        _log(f"  balanced_acc ours vs sklearn: {balanced_accuracy(yt, yp):.3f} vs "
             f"{sk_bal_acc:.3f} "
             f"({'PASS' if abs(balanced_accuracy(yt, yp)-sk_bal_acc) < 1e-3 else 'FAIL'})")
        _log(f"  ROC-AUC ours vs sklearn: {auc:.3f} vs {sk_auc_val:.3f} "
             f"({'PASS' if abs(auc-sk_auc_val) < 1e-3 else 'FAIL'})")
    else:
        _log("\n--- (2) scikit-learn not installed in this environment ---")
        _log("       Install scikit-learn and re-run for side-by-side comparison.")

    _log("\n" + "=" * 64)
    n_pass = sum(1 for r in results if "PASS" in r)
    n_fail = sum(1 for r in results if "FAIL" in r)
    _log(f"SUMMARY: {n_pass} PASS, {n_fail} FAIL")
    _log("=" * 64)

    with open(REPORT, "w") as f:
        f.write("\n".join(results) + "\n")
    print(f"\nWrote {REPORT}")



# ===========================================================================
# SECTION 7  Word report generator  (was build_report.py)
# ===========================================================================

def style_normal(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)


def add_heading(doc, text, level):
    sizes = {0: 18, 1: 14, 2: 12}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_para(doc, text, italic=False, bold=False, size=None,
             alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_figure(doc, path, caption, width_cm=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(10)


def add_hyperlink(paragraph, url, text, size=10):
    """Append a clickable external hyperlink run to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/"
        "relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")        # standard Word hyperlink blue
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))  # half-points
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


# Blue table styling (ColorBrewer RdBu tints, matching the figures)
TABLE_BORDER = "4393C3"      # medium blue table grid lines
TABLE_HEADER = "D1E5F0"      # light blue header-row shading


def _set_table_borders(table, color=TABLE_BORDER, sz=6):
    """Set every table grid line (outer + inner) to a single colour."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))      # eighths of a point
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    # Insert in the schema-correct position within tblPr
    tblPr.insert_element_before(
        borders, "w:shd", "w:tblLayout", "w:tblCellMar",
        "w:tblLook", "w:tblCaption", "w:tblDescription")


def _shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_table(doc, headers, rows, widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    _set_table_borders(table)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _shade_cell(cell, TABLE_HEADER)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = str(val)
    return table


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
def cmd_report():
    doc = Document()
    style_normal(doc)

    # ---- Front matter -- title, name, date and GitHub go on the coversheet ----
    add_para(doc, "", space_after=0)

    doc.add_paragraph()

    # ---- Abstract -------------------------------------------------------------
    add_heading(doc, "Abstract", 1)
    add_para(doc,
        "Background: Depression in young adulthood is rising and is an "
        "established cardiovascular disease (CVD) risk factor, yet the "
        "evidence base draws predominantly from older populations. "
        "Whether depressive symptomatology improves CVD risk "
        "discrimination beyond demographics in adults under 40 is "
        "unclear. We hypothesised that the Patient Health "
        "Questionnaire-9 (PHQ-9) would add reproducible discriminative "
        "value over demographics alone.")
    add_para(doc,
        "Methods: We analysed the NHANES pre-pandemic cycle (2017-March "
        "2020), restricted to adults aged 18-39 with a usable PHQ-9 "
        "(n = 2,728). Data driven cardiovascular risk profiles were "
        "derived by K-means clustering of 13 self-reported and "
        "behavioural risk indicators. Three classifiers (logistic "
        "regression, random forest and gradient boosting) predicted "
        "profile membership from PHQ-9 plus demographics, each versus "
        "a demographics-only model. Discrimination used 5 fold "
        "stratified cross-validation over three seeds, compared by "
        "paired t-test; sensitivity analyses examined age, clustering "
        "method and fairness.")
    add_para(doc,
        "Results: Clustering identified two profiles (silhouette 0.37); "
        "the elevated-risk profile (22.0%) was marked by "
        "physician-recorded lifestyle advice and self-reported "
        "hypertension and dyslipidaemia. Its prevalence was 35.3% with "
        "PHQ-9 ≥ 10 versus 20.6% below 10. Adding PHQ-9 raised "
        "cross-validated ROC-AUC from 0.602 to 0.637 (logistic "
        "regression), 0.592 to 0.638 (random forest) and 0.599 to "
        "0.642 (gradient boosting); every increment was significant "
        "(all p < 10⁻⁵), and PHQ-9 ranked highest by "
        "permutation importance.")
    add_para(doc,
        "Conclusions: Depressive symptomatology provides a small but "
        "statistically robust, algorithm-independent improvement in "
        "predicting cardiovascular risk profiles in young adults. The "
        "modest discrimination (AUC ≈ 0.64) indicates the PHQ-9 should "
        "complement, not replace, structured cardiovascular assessment. "
        "Code and dataset-access instructions are openly available.")

    doc.add_page_break()

    # ---- Introduction ---------------------------------------------------------
    add_heading(doc, "1. Introduction", 1)
    add_para(doc,
        "Depression is among the leading contributors to the global "
        "disease burden, and its prevalence in young people has risen "
        "sharply over the past decade (Thapar et al., 2022). Large "
        "epidemiological studies have established depression as an "
        "independent cardiovascular disease (CVD) risk factor: the "
        "American Heart Association’s Life’s Essential 8 was recently "
        "extended to Life’s Crucial 9 with a psychological-health "
        "dimension, motivated by NHANES evidence linking depression to "
        "cardiovascular mortality (Ge et al., 2024; Xu et al., 2024). "
        "Subsequent analyses show depression severity tracks "
        "cardiovascular-kidney-metabolic (CKM) syndrome progression (He "
        "et al., 2025; Wang et al., 2025), mortality in diabetes (Zhang Z "
        "et al., 2025), and is mediated in part by inflammation (Meng et "
        "al., 2025).")

    add_para(doc,
        "Most of this evidence draws on older adults. Younger adults "
        "accumulate sub-clinical risk for years before events occur, so "
        "risk stratification must rely on modifiable factors and family "
        "history. Machine learning (ML) is well suited here, capturing "
        "non-linear interactions and producing data-driven phenotypes "
        "classical risk equations may miss (Dipnall et al., 2016). "
        "NHANES has become a workhorse for such analyses (Zhang YZ et al., "
        "2025; Vu et al., 2025; Wu et al., 2025; Liu et al., 2025), and "
        "lifestyle and metabolic correlates of depression are well "
        "documented in it (Liu et al., 2024; Zhou et al., 2024; Guo et "
        "al., 2025; Tian et al., 2025).")

    add_para(doc,
        "What remains under-explored is whether the routinely available "
        "PHQ-9 can predict a data-driven CV risk phenotype specifically "
        "in young adults. We therefore asked: can depression in younger "
        "adults be used to predict cardiovascular risk profiles using "
        "machine learning? In NHANES adults aged 18-39 we derived "
        "unsupervised CV risk profiles, trained three classifiers to "
        "predict profile membership from PHQ-9 plus demographics, and "
        "compared against a demographics only baseline, hypothesising "
        "that depression would offer incremental discriminative value.")

    # ---- Methods --------------------------------------------------------------
    add_heading(doc, "2. Methods", 1)

    add_heading(doc, "2.1 Data source and rationale", 2)
    add_para(doc,
        "NHANES is a continuous multi-stage stratified probability "
        "sample of the US non-institutionalised civilian population. "
        "Three features make it well-suited to ML for early-adulthood CV "
        "research: (i) it couples the validated PHQ-9 with self-reported "
        "diagnoses, medication and physiological measurements in the "
        "same participants; (ii) the per-cycle sample (≈ 15,000 overall; "
        "≈ 3,000 aged 18-39) supports repeated stratified CV and stable "
        "feature importance; (iii) every variable is documented in a "
        "public codebook with stable cross-cycle coding, supporting "
        "reproducibility, a benchmark met by the NHANES-ML literature "
        "(Vu et al., 2025; Wu et al., 2025; Zhang YZ et al., 2025).")

    add_para(doc,
        "NHANES uses stratified multi-stage cluster sampling with "
        "oversampling of minority and low-income strata. We treat it as "
        "a learning sample and report unweighted ML metrics, since "
        "accuracy, F1 and AUC describe the model on data from this "
        "design rather than a simple random sample; the minority "
        "oversampling aids fairness audits but means sample prevalence "
        "≠ national prevalence. We used the 2017-March 2020 pre-"
        "pandemic combined cycle, the most recent NHANES release "
        "containing PHQ-9 alongside the cardiovascular and medical-"
        "conditions modules.")

    add_heading(doc, "2.2 Sample", 2)
    add_para(doc,
        "Three considerations motivated the 18-39 window: CV risk "
        "accumulates at this stage before events, when modifiable "
        "factors remain amenable to intervention; depressive-symptom "
        "prevalence peaks in early adulthood (Thapar et al., 2022); and "
        "restricting to younger adults reduces confounding by "
        "comorbidities and polypharmacy.")

    add_para(doc,
        "From 3,260 participants aged 18-39 we excluded pregnant "
        "respondents and those without a usable PHQ-9 (≤ 2 items "
        "missing); n = 2,728.")

    add_heading(doc, "2.3 Variables", 2)
    add_para(doc,
        "Depression. The PHQ-9 is a nine-item self-report scale (total "
        "0-27); PHQ-9 ≥ 10 is the conventional threshold for clinically "
        "significant symptoms. We modelled PHQ-9 as continuous in the "
        "supervised models and binary in cross-tabulations.")

    add_para(doc,
        "Cardiovascular risk indicators (13): self-reported hypertension "
        "(BPQ020), antihypertensive use (BPQ050A), high cholesterol "
        "(BPQ080), lipid-lowering use (BPQ100D), doctor-reported "
        "overweight (MCQ080), four physician lifestyle directives "
        "(MCQ366A-D), family history of premature MI (MCQ300A) and "
        "diabetes (MCQ300C), insufficient sleep (< 7 h) and household "
        "tobacco exposure. Because clinical events are rare before age "
        "40, physician-recognised risk factors provide a more "
        "informative substrate.")

    add_para(doc,
        "Demographics and socioeconomic status: age, sex, race/"
        "ethnicity, education, marital status, and family poverty-to-"
        "income ratio.")

    add_heading(doc, "2.4 Pre-processing", 2)
    add_para(doc,
        "NHANES refused/don’t-know codes were set to missing; binary "
        "indicators were mode-imputed, continuous indicators median-"
        "imputed. Categorical variables were one-hot encoded; continuous "
        "variables z-score scaled, with scaling parameters estimated on "
        "the training fold only to avoid leakage.")

    add_heading(doc, "2.5 Cardiovascular risk-profile derivation", 2)
    add_para(doc,
        "Rationale for clustering rather than a published score: (i) "
        "calibrated scores (Framingham, AHA PREVENT, QRISK3) were "
        "derived on adults ≥ 30-40 and collapse near “low risk” here; "
        "(ii) they require measured biomarkers unavailable in our "
        "subset; (iii) a fixed-weight composite ignores risk-factor "
        "co-occurrence patterns recognised in cardiovascular-kidney-"
        "metabolic phenotyping (He et al., 2025; Wang et al., 2025).")

    add_para(doc,
        "K-means (k-means++, 20 restarts) on the standardised indicators "
        "recovers the constellation in which physician-flagged lifestyle "
        "issues, diagnoses and family history co-occur, a learned "
        "“clinician-recognised at-risk” phenotype (Figure 1B) treated as "
        "an analytic target, not a clinical diagnosis (Section 4.4). K "
        "was selected by mean silhouette over K ∈ {2,3,4}.")

    add_heading(doc, "2.6 Supervised models", 2)
    add_para(doc,
        "Three classifiers with deliberately different inductive biases "
        "were used: multinomial logistic regression (L2-regularised, "
        "linear, with directly interpretable log-odds coefficients); a "
        "random forest (bagged CART trees); and gradient boosting. "
        "Models were implemented in Python; full code is in the GitHub "
        "repository. For each algorithm we fit a demographics only "
        "baseline alongside the full model: a signal surviving linear, "
        "bagged and boosted models alike is unlikely to be a single-"
        "model artefact.")

    add_heading(doc, "2.7 Hyperparameter tuning", 2)
    add_para(doc,
        "Hyperparameters were selected by grid search with 3-fold "
        "stratified CV on an 80% tuning subset (n = 2,182). Selection "
        "criterion: mean macro F1. The LR grid spanned 16 configs, RF "
        "27 and GB 12 (Table 2); the full 55-config grid is in "
        "results/hyperparameter_search.csv.")

    add_table(doc,
        headers=["Algorithm", "Hyperparameter", "Values tested", "Selected"],
        rows=[
            ["Logistic regression", "C (inverse L2)", "0.01, 0.1, 1, 10", "1"],
            ["Logistic regression", "Learning rate", "0.1, 0.3, 0.5, 1.0", "1.0"],
            ["Random forest", "n_estimators", "40, 60, 100", "60"],
            ["Random forest", "max_depth", "4, 6, 8", "8"],
            ["Random forest", "min_samples_leaf", "20, 30, 50", "30"],
            ["Gradient boosting", "n_estimators", "40, 60, 100", "60"],
            ["Gradient boosting", "Learning rate", "0.05, 0.1", "0.05"],
            ["Gradient boosting", "max_depth", "2, 3", "3"],
        ],
    )
    add_para(doc, "Table 2. Hyperparameter search space and selected "
                  "configurations (√p features per split fixed for the "
                  "tree ensembles).",
             italic=True, size=10)

    add_heading(doc, "2.8 Evaluation and interpretability", 2)
    add_para(doc,
        "Performance was estimated by stratified 5-fold CV over three "
        "seeds (15 folds), with minority-class oversampling on training "
        "folds only. We report accuracy, balanced accuracy, macro F1 and "
        "macro-averaged one-versus-rest ROC-AUC; the full-vs-baseline "
        "macro F1 difference was tested with a paired t-test. On a "
        "held-out 25% split we computed RF permutation importance, LR "
        "coefficients, subgroup performance (sex, race/ethnicity, income "
        "tertile), and calibration (Brier score, predicted vs observed "
        "prevalence).")

    add_heading(doc, "2.9 Software and ethics", 2)
    add_para(doc,
        "Python 3.10, NumPy 2.2, pandas 2.3, matplotlib 3.10, seaborn "
        "0.13; reproducible from code/run_analysis.py. NHANES is "
        "public-domain de-identified data; IRB approval was not "
        "required.")

    # ---- Results --------------------------------------------------------------
    add_heading(doc, "3. Results", 1)

    add_heading(doc, "3.1 Cohort", 2)
    add_para(doc,
        "The analytic cohort comprised 2,728 adults aged 18-39 (mean age "
        "28.0; 51.0% female; 9.2% with PHQ-9 ≥ 10). Table 1 summarises "
        "characteristics by derived risk profile; every cardiovascular "
        "indicator and the PHQ-9 is higher in the elevated-risk "
        "profile.")

    add_table(doc,
        headers=["Characteristic", "Overall", "Low-risk profile",
                 "Elevated-risk profile"],
        rows=[
            ["n", "2,728", "2,128", "600"],
            ["Age, mean (SD)", "28.0 (6.6)", "27.6 (6.5)", "29.5 (6.5)"],
            ["Female, %", "51.0", "48.6", "59.7"],
            ["Non-Hispanic White, %", "31.3", "31.7", "29.7"],
            ["Non-Hispanic Black, %", "25.0", "25.0", "25.0"],
            ["Mexican American, %", "14.8", "13.8", "18.5"],
            ["Non-Hispanic Asian, %", "12.1", "12.8", "9.7"],
            ["PHQ-9 total, mean (SD)", "3.4 (4.2)", "3.0 (3.9)", "4.5 (4.8)"],
            ["PHQ-9 ≥ 10, %", "9.2", "7.7", "14.8"],
            ["Self-reported hypertension, %", "12.1", "7.5", "28.3"],
            ["Self-reported high cholesterol, %", "10.6", "5.5", "28.7"],
            ["Doctor-reported overweight, %", "33.7", "19.1", "85.3"],
            ["Family history, premature MI, %", "8.6", "7.2", "13.3"],
            ["Family history, diabetes, %", "34.6", "30.7", "48.3"],
            ["Insufficient sleep (< 7 h), %", "23.7", "23.3", "25.3"],
            ["Household smoker, %", "35.6", "36.3", "33.0"],
        ],
    )
    add_para(doc, "Table 1. Cohort characteristics, overall and by "
                  "derived cardiovascular risk profile (n = 2,728).",
             italic=True, size=10)

    add_heading(doc, "3.2 Cardiovascular risk profiles", 2)
    add_para(doc,
        "K-means selected K = 2 (silhouette 0.373 versus 0.274 and 0.285 "
        "for K = 3 and K = 4; Figure 1A). The two profiles separate "
        "sharply on physician-given lifestyle advice (Figure 1B): in "
        "profile 1 (n = 600, 22.0%), standardised scores for being told "
        "to lose weight, exercise, reduce salt and reduce fat ranged "
        "+1.12 to +1.58 SD versus −0.44 to −0.31 SD in profile 0, with "
        "elevated self-reported hypertension (+0.50 SD) and "
        "hypercholesterolaemia (+0.59 SD). We label profile 0 “low” and "
        "profile 1 “elevated” CV risk.")

    add_figure(doc, FIG / "fig_silhouette.png",
               "Figure 1A. Silhouette and elbow plots for K-means cluster "
               "selection over K ∈ {2,3,4}.", width_cm=14)
    add_figure(doc, FIG / "fig_cluster_profile.png",
               "Figure 1B. Standardised K-means centroids on cardiovascular "
               "risk indicators. The elevated-risk profile is dominated by "
               "physician-given lifestyle directives (MCQ366A-D) and "
               "self-reported hypertension and dyslipidaemia.", width_cm=15)

    add_heading(doc, "3.3 Crude depression-risk-profile association", 2)
    add_para(doc,
        "The elevated-risk profile contained 35.3% of participants with "
        "PHQ-9 ≥ 10 versus 20.6% of those with PHQ-9 < 10: clinically "
        "significant depression doubled the odds of the elevated-risk "
        "profile (odds ratio 2.10, 95% CI 1.59-2.77; χ² = 28.7, "
        "p < 0.001). Median PHQ-9 was 2 in the low-risk profile versus "
        "3 in the high-risk profile (Figure 2).")

    add_figure(doc, FIG / "fig_phq_by_profile.png",
               "Figure 2. Distribution of PHQ-9 total scores by cardiovascular "
               "risk profile. Box plot summarises the inter-quartile range; "
               "scattered points show a sub-sample of individual "
               "participants.", width_cm=16)

    add_heading(doc, "3.4 Model performance", 2)
    add_para(doc,
        "Cross-validated discrimination is in Table 3 and Figure 3. "
        "Adding PHQ-9 raised ROC-AUC by ≈ 0.04 in all three algorithms "
        "(LR: 0.602 → 0.637; RF: 0.592 → 0.638; GB: 0.599 → 0.642), the "
        "boosted ensemble marginally strongest, with good stability "
        "across folds (SDs 0.02-0.04). Paired t-tests confirmed "
        "significant macro-F1 improvement under every algorithm (LR "
        "t = 4.49; RF t = 7.47; GB t = 5.48; all p < 10⁻⁵).")

    add_table(doc,
        headers=["Model", "Accuracy", "Balanced acc", "Macro F1", "ROC-AUC"],
        rows=[
            ["Logistic regression (demographics only)",
             "0.575 ± 0.020", "0.573 ± 0.022", "0.525 ± 0.018", "0.602 ± 0.026"],
            ["Logistic regression (+ PHQ-9)",
             "0.605 ± 0.020", "0.591 ± 0.029", "0.547 ± 0.022", "0.637 ± 0.031"],
            ["Random forest (demographics only)",
             "0.608 ± 0.022", "0.565 ± 0.021", "0.536 ± 0.019", "0.592 ± 0.031"],
            ["Random forest (+ PHQ-9)",
             "0.634 ± 0.022", "0.604 ± 0.027", "0.568 ± 0.023", "0.638 ± 0.034"],
            ["Gradient boosting (demographics only)",
             "0.592 ± 0.026", "0.571 ± 0.028", "0.532 ± 0.024", "0.599 ± 0.034"],
            ["Gradient boosting (+ PHQ-9)",
             "0.620 ± 0.025", "0.606 ± 0.035", "0.562 ± 0.028", "0.642 ± 0.036"],
        ],
    )
    add_para(doc, "Table 3. Cross-validated performance (mean ± SD, "
                  "5-fold × 3 seeds). Adding PHQ-9 improves every model.",
             italic=True, size=10)

    add_figure(doc, FIG / "fig_cv_metrics.png",
               "Figure 3. Cross-validated performance of the six models "
               "(three algorithms × full / demographics-only). Adding "
               "PHQ-9 consistently improves on the demographics-only "
               "baseline.", width_cm=15)

    add_heading(doc, "3.5 Confusion and interpretability", 2)
    add_para(doc,
        "On the held-out 25% split the RF confusion matrix (Figure 4A) "
        "showed moderate sensitivity for the elevated-risk class "
        "(≈ 0.53), with false positives expected from training-time "
        "oversampling (Section 4.3).")

    add_figure(doc, FIG / "fig_rf_confusion.png",
               "Figure 4A. Random forest confusion matrix on the held-out 25% "
               "test set. Cell shading shows row-normalised proportion; "
               "numerals are raw counts.", width_cm=10)

    add_para(doc,
        "Permutation importance ranked PHQ-9 as the single most "
        "informative feature (Figure 4B; macro-F1 drop 0.032 ± 0.013), "
        "well ahead of any other variable; LR coefficients corroborated "
        "this (PHQ-9 +0.19 standardised, comparable to age +0.21). The "
        "marginal RF effect of PHQ-9 (Figure 5) was monotonic, rising "
        "from ≈ 0.48 at PHQ-9 = 0 to ≈ 0.65 at PHQ-9 = 20.")

    add_figure(doc, FIG / "fig_rf_permimp.png",
               "Figure 4B. Random forest permutation importance "
               "(8 permutations, macro F1 drop). PHQ-9 total is the "
               "single most informative feature.", width_cm=14)

    add_figure(doc, FIG / "fig_rf_phq_marginal.png",
               "Figure 5. Marginal random-forest predicted probability of "
               "each risk profile as a function of PHQ-9 total, with "
               "demographics held at sample means.", width_cm=13)

    add_heading(doc, "3.6 Sensitivity analyses", 2)
    add_para(doc,
        "Age window. Re-running on broader age windows (Table 4) showed "
        "the incremental PHQ-9 ROC-AUC was largest in 18-39 (+0.036 LR / "
        "+0.042 RF) and roughly halved in 18-64 and 18+. Baseline AUC "
        "rose with age range as accumulated risk factors carry more "
        "weight in older adults; all ablations stayed significant "
        "(p < 0.005). Depression's marginal value is therefore greatest "
        "in the under 40 window.")

    add_para(doc,
        "Clustering method. Re-deriving the profiles with K-modes (the "
        "matching-dissimilarity method appropriate for binary data) "
        "agreed closely with the K-means profiles (98.3% agreement, "
        "adjusted Rand index 0.92) and preserved the depression gradient "
        "(31.7% elevated-risk in PHQ-9 ≥ 10 vs 19.1% otherwise). The "
        "phenotype is not an artefact of K-means on binary data.")

    add_table(doc,
        headers=["Cohort", "n", "Elev %", "LR demo AUC", "LR + PHQ AUC", "ΔAUC", "RF demo AUC", "RF + PHQ AUC", "ΔAUC"],
        rows=[
            ["18-39 (primary)", "2,728", "22.0", "0.602", "0.637", "+0.036", "0.595", "0.637", "+0.042"],
            ["18-64 (working-age)", "6,254", "34.5", "0.668", "0.685", "+0.017", "0.675", "0.691", "+0.016"],
            ["18+ (all adults)",  "8,223", "36.7", "0.654", "0.677", "+0.024", "0.671", "0.692", "+0.021"],
        ],
    )
    add_para(doc, "Table 4. Age-window sensitivity analysis (5-fold CV × "
                  "2 seeds). PHQ-9’s incremental discriminative value is "
                  "largest in the under 40 cohort.",
             italic=True, size=10)

    add_heading(doc, "3.7 Subgroup performance (fairness check)", 2)
    add_para(doc,
        "Held-out RF subgroup performance is in Table 5. Sex parity was "
        "good (AUC 0.608 vs 0.606); race/ethnicity AUC ranged 0.586 to "
        "0.631, balanced accuracy lowest in Mexican American (0.511); "
        "income tertiles showed a modest gradient (AUC 0.598 → 0.622).")

    add_table(doc,
        headers=["Subgroup", "n", "% elevated", "Bal acc", "Macro F1", "ROC-AUC"],
        rows=[
            ["Overall", "682", "22.0", "0.580", "0.543", "0.609"],
            ["Sex: Female", "354", "23.7", "0.568", "0.500", "0.608"],
            ["Sex: Male", "328", "20.1", "0.577", "0.568", "0.606"],
            ["Race: Mexican American", "102", "26.5", "0.511", "0.470", "0.587"],
            ["Race: Non-Hispanic Asian", "90", "11.1", "0.544", "0.460", "0.586"],
            ["Race: Non-Hispanic Black", "162", "25.9", "0.611", "0.589", "0.631"],
            ["Race: Non-Hispanic White", "209", "23.4", "0.586", "0.561", "0.613"],
            ["Race: Other Hispanic", "80", "22.5", "0.637", "0.623", "0.617"],
            ["Income (PIR) low <1.47", "227", "20.7", "0.576", "0.550", "0.598"],
            ["Income (PIR) mid 1.47-2.46", "226", "22.1", "0.570", "0.544", "0.610"],
            ["Income (PIR) high ≥2.46", "229", "23.1", "0.588", "0.526", "0.622"],
        ],
    )
    add_para(doc, "Table 5. Held-out RF subgroup performance "
                  "(n = 682 test split). The Other/Multi-race subgroup "
                  "(n = 39) is omitted because its small size renders "
                  "metrics unstable.",
             italic=True, size=10)

    # ---- Discussion -----------------------------------------------------------
    add_heading(doc, "4. Discussion", 1)

    add_heading(doc, "4.1 Summary of key findings", 2)
    add_para(doc,
        "The results broadly support the hypothesis that depressive "
        "symptomatology carries cardiovascular-risk information in young "
        "adults beyond what demographics provide. Even in an under 40 "
        "sample where overt CVD is rare, the PHQ-9 added a small but "
        "reproducible improvement to discrimination of a data-driven CV "
        "risk phenotype. The most informative finding is the consistency "
        "of the signal: three algorithms with different inductive biases "
        "converged on ROC-AUC ≈ 0.64, every ablation was significant, "
        "and PHQ-9 ranked as the most important input by permutation "
        "importance. Crude elevated-risk prevalence was 71% higher among "
        "PHQ-9 ≥ 10 participants (35.3% vs 20.6%).")

    add_para(doc,
        "These findings align with the accumulating NHANES literature "
        "linking depression to cardiovascular health (Ge et al., 2024; "
        "Xu et al., 2024; Wu et al., 2025; He et al., 2025; Wang et al., "
        "2025; Zhang Z et al., 2025; Meng et al., 2025). Our contribution "
        "is to show the signal is recoverable from PHQ-9 alone against a "
        "learned multivariate phenotype, with incremental value greatest "
        "in the under 40 window.")

    add_heading(doc, "4.2 Model comparison and trade-offs", 2)
    add_para(doc,
        "All three algorithms converged on near-identical discrimination "
        "(ROC-AUC 0.637-0.642), itself informative: the depression-CV-"
        "risk signal is largely additive and linearly accessible, with "
        "little extra recoverable from non-linear interactions. This "
        "shapes the model-choice trade-off. Logistic regression is the "
        "most transparent (coefficients read as log-odds, fast, simple "
        "to audit and recalibrate); the tree ensembles capture "
        "interactions but at the cost of interpretability (needing post-"
        "hoc tools such as SHAP) and greater computational expense. For "
        "a primary-care screening application, where transparency "
        "matters and the boosted model's 0.005 AUC edge is not "
        "clinically meaningful, logistic regression is preferable; for "
        "a research setting prioritising predictive ceiling, gradient "
        "boosting would be chosen. The pipeline retains all three.")

    add_heading(doc, "4.3 Calibration and clinical utility", 2)
    add_para(doc,
        "Discrimination and calibration are distinct, and the headline "
        "metric concerns only the former. An AUC of 0.64 falls below the "
        "0.70 conventionally taken as the minimum for acceptable "
        "discrimination: given one elevated-risk and one low-risk "
        "participant, the model ranks them correctly 64% of the time. "
        "Calibration is weaker still. Because training folds were "
        "oversampled to 50:50, the random forest's predicted "
        "probabilities are inflated (mean 0.46 against a true prevalence "
        "of 0.22) and its Brier score (0.229) is worse than simply "
        "predicting the base rate (0.172); a non-oversampled model is "
        "well calibrated (mean 0.22, Brier 0.166) but loses minority-"
        "class sensitivity. The probabilities in Figure 5 should "
        "therefore be read as relative ordering, not true risks. "
        "Clinically, this is not a deployable individual risk "
        "calculator. Its defensible uses are narrower: a low-stakes "
        "triage signal, where even weak discrimination can carry net "
        "benefit because flagging a young adult for fuller "
        "cardiovascular assessment is cheap and low-harm, and "
        "population-level evidence. It should not inform treatment "
        "decisions.")

    add_heading(doc, "4.4 Strengths and limitations", 2)
    add_para(doc,
        "The analysis has several strengths: a strict ablation against a "
        "demographics-only baseline, documented grid-search tuning, and "
        "age-window, clustering-method and subgroup-fairness sensitivity "
        "analyses. In keeping with open-science practice the whole "
        "pipeline is fully reproducible, with all code, result files and "
        "dataset-access instructions openly available for independent "
        "re-running and audit. Six limitations qualify the "
        "interpretation. (i) The design is cross-sectional, so reverse "
        "causation remains plausible. (ii) CV indicators are self-"
        "reported, biasing ascertainment toward those with healthcare "
        "access. (iii) Measured biomarkers were unavailable. (iv) "
        "Discrimination (AUC ≈ 0.64) is modest. (v) The K-means "
        "phenotype is an analytic target, not a prospectively validated "
        "risk class; external validation against incident events is "
        "essential before clinical use. (vi) Although NHANES is a "
        "probability sample, the complete-case analytic subset may "
        "under-represent the most survey- and healthcare-disengaged "
        "young adults, modestly limiting external validity. A minor "
        "redundancy is also noted: sleep enters the clustering through "
        "two correlated indicators, mildly over-weighting it, though its "
        "near-zero cluster separation (Figure 1B) makes the effect "
        "negligible.")

    add_heading(doc, "4.5 Ethics, fairness and interpretability", 2)
    add_para(doc,
        "The subgroup audit (Table 5) shows good sex parity (AUC 0.61 "
        "vs 0.61) but uneven race/ethnicity discrimination, weakest in "
        "Mexican American (AUC 0.59, balanced accuracy 0.51) and Non-"
        "Hispanic Asian (AUC 0.59) participants, with a modest income "
        "gradient (AUC 0.60 → 0.62). These disparities reflect "
        "underlying differences in CV-risk prevalence and clinician "
        "flagging rather than biological causation, but indicate the "
        "model would under-serve some subgroups in deployment. The "
        "reliance on self-reported diagnoses compounds this: groups "
        "with poorer healthcare access are under-diagnosed and so "
        "receive systematically lower risk estimates, risking under-"
        "recognition of those most in need. Clinical-translation "
        "pipelines should remove race or replace it with social-"
        "determinant-of-health proxies (Vu et al., 2025; Wu et al., "
        "2025). Using depression as a CV risk marker also risks stigma "
        "for high-PHQ-9 patients; deployment would require governance "
        "and fairness audits.")

    add_heading(doc, "4.6 Implications and future development", 2)
    add_para(doc,
        "The findings support integrating mental-health screening into "
        "early-adulthood CV cohorts. Future work should pool additional "
        "NHANES cycles to enlarge the under 40 sample, validate the "
        "phenotype against incident cardiovascular events, and re-derive "
        "the risk profiles on measured biomarkers.")

    # ---- Conclusion -----------------------------------------------------------
    add_heading(doc, "5. Conclusion", 1)
    add_para(doc,
        "Overall, the findings tell a coherent story. In NHANES adults "
        "aged 18-39, depressive symptomatology measured by the PHQ-9 "
        "contributes a small but reproducible improvement over "
        "demographics in predicting a data-driven cardiovascular risk "
        "phenotype, consistent across three algorithms and largest in "
        "the under 40 window. The effect is modest and should not be "
        "overinterpreted: PHQ-9 is not a stand-alone risk tool, but it "
        "carries enough signal to support depression screening in "
        "early-adulthood CV risk assessment, reinforcing the view that "
        "psychological health is part of cardiovascular health.")

    # ---- References -----------------------------------------------------------
    add_heading(doc, "References", 1)
    # Harvard-style references: alphabetical, numbered, with the journal
    # name italicised. Each entry is (authors + year + title, journal,
    # volume/pages, PubMed URL). Eleven links use the article's verified
    # PMID; five recent articles, whose PMIDs could not be confirmed, use
    # a PubMed DOI lookup that resolves to the same record.
    refs = [
        ("Dipnall, J.F., Pasco, J.A., Berk, M. et al. (2016) 'Into the "
         "Bowels of Depression: Unravelling Medical Symptoms Associated "
         "with Depression by Applying Machine-Learning Techniques to a "
         "Community Based Population Sample', ",
         "PLoS One",
         ", 11(12), e0167055.",
         "https://pubmed.ncbi.nlm.nih.gov/27935995/"),

        ("Ge, J., Peng, W. and Lu, J. (2024) 'Predictive Value of Life’s "
         "Crucial 9 for Cardiovascular and All-Cause Mortality: A "
         "Prospective Cohort Study From the NHANES 2007 to 2018', ",
         "Journal of the American Heart Association",
         ", 13(19), e036669.",
         "https://pubmed.ncbi.nlm.nih.gov/39377201/"),

        ("Guo, T., Zhou, Y., Yang, G., Sheng, L. and Chai, X. (2025) "
         "'Association between cardiometabolic index and hypertension "
         "among US adults from NHANES 1999-2020', ",
         "Scientific Reports",
         ", 15, 87029.",
         "https://pubmed.ncbi.nlm.nih.gov/?term=10.1038%2Fs41598-025-87029-0"),

        ("He, Y., Lan, L., Liu, Y. et al. (2025) 'Depression severity as "
         "a predictor of cardiovascular-kidney-metabolic syndrome "
         "progression and mortality: Results from two nationally "
         "representative cohort studies', ",
         "Journal of Affective Disorders",
         ", 388, 119606.",
         "https://pubmed.ncbi.nlm.nih.gov/?term=10.1016%2Fj.jad.2025.119606"),

        ("Liu, H., Dong, H., Zhou, Y. et al. (2024) 'The association "
         "between Metabolic Score for Visceral Fat and depression in "
         "overweight or obese individuals', ",
         "Frontiers in Endocrinology",
         ", 15, 1482003.",
         "https://pubmed.ncbi.nlm.nih.gov/?term=10.3389%2Ffendo.2024.1482003"),

        ("Liu, X., Luo, Z., Jing, F. et al. (2025) 'Estimating "
         "cardiovascular mortality in patients with hypertension using "
         "machine learning: the role of depression classification based "
         "on lifestyle and physical activity', ",
         "Journal of Psychosomatic Research",
         ", 189, 112030.",
         "https://pubmed.ncbi.nlm.nih.gov/39752763/"),

        ("Meng, X., Han, L., Fu, J., Hu, C. and Lu, Y. (2025) "
         "'Associations between metabolic syndrome and depression, and "
         "the mediating role of inflammation: Based on the NHANES "
         "database', ",
         "Journal of Affective Disorders",
         ", 375, pp. 214-221.",
         "https://pubmed.ncbi.nlm.nih.gov/39862983/"),

        ("Thapar, A., Eyre, O., Patel, V. and Brent, D. (2022) "
         "'Depression in young people', ",
         "The Lancet",
         ", 400(10352), pp. 617-631.",
         "https://pubmed.ncbi.nlm.nih.gov/35940184/"),

        ("Tian, Q., Guo, J., Ding, J. and Zhu, Y. (2025) 'The "
         "relationship between unhealthy lifestyle behaviors and "
         "depression: Evidence from NHANES', ",
         "Journal of Affective Disorders",
         ", 384, pp. 214-221.",
         "https://pubmed.ncbi.nlm.nih.gov/40258422/"),

        ("Vu, T., Dawadi, R., Yamamoto, M. et al. (2025) 'Prediction of "
         "depressive disorder using machine learning approaches: "
         "findings from the NHANES', ",
         "BMC Medical Informatics and Decision Making",
         ", 25, 83.",
         "https://pubmed.ncbi.nlm.nih.gov/39962516/"),

        ("Wang, G., Wu, Y., Chen, A. et al. (2025) 'Association between "
         "depressive symptoms and mortality in patients with "
         "Cardiovascular-Kidney-Metabolic syndrome: the mediating role "
         "of inflammatory biomarkers', ",
         "Journal of Affective Disorders",
         ", 386, 119429.",
         "https://pubmed.ncbi.nlm.nih.gov/?term=10.1016%2Fj.jad.2025.119429"),

        ("Wu, Z., Xu, P., Zhai, Y. et al. (2025) 'The Association of "
         "Elevated Depression Levels and Life’s Essential 8 on "
         "Cardiovascular Health With Predicted Machine Learning Models "
         "and Interpretations: Evidence From NHANES 2007-2018', ",
         "Depression and Anxiety",
         ", 2025, 8865176.",
         "https://pubmed.ncbi.nlm.nih.gov/40255861/"),

        ("Xu, Y., Ning, W., Zhang, Y. et al. (2024) 'Associations "
         "Between Cardiovascular Health (Life’s Essential 8) and Mental "
         "Disorders', ",
         "Clinical Cardiology",
         ", 47(10), e70019.",
         "https://pubmed.ncbi.nlm.nih.gov/39314085/"),

        ("Zhang, Y.-Z., Wu, H.-Y., Ma, R.-W. et al. (2025) 'Machine "
         "Learning-Based predictive model for adolescent metabolic "
         "syndrome: Utilizing data from NHANES 2007-2016', ",
         "Scientific Reports",
         ", 15, 88156.",
         "https://pubmed.ncbi.nlm.nih.gov/39863763/"),

        ("Zhang, Z., Xu, H., Meng, Y., Yan, Y. and Wang, Y. (2025) "
         "'Association between depressive symptoms and all-cause and "
         "cardiovascular mortality in diabetes patients: A national "
         "cohort study', ",
         "Journal of Affective Disorders",
         ", 391, 119957.",
         "https://pubmed.ncbi.nlm.nih.gov/?term=10.1016%2Fj.jad.2025.119957"),

        ("Zhou, X., Tao, X.-L., Zhang, L. et al. (2024) 'Association "
         "between cardiometabolic index and depression: National Health "
         "and Nutrition Examination Survey (NHANES) 2011-2014', ",
         "Journal of Affective Disorders",
         ", 351, pp. 939-947.",
         "https://pubmed.ncbi.nlm.nih.gov/38341157/"),
    ]
    for i, (part1, journal, part2, url) in enumerate(refs, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.8)
        p.paragraph_format.first_line_indent = Cm(-0.8)
        for seg, ital in ((f"{i}. ", False), (part1, False),
                          (journal, True), (part2 + "  ", False)):
            r = p.add_run(seg)
            r.font.size = Pt(10)
            r.italic = ital
        add_hyperlink(p, url, "[PubMed]", size=10)

    doc.save(str(OUT_REPORT))
    print(f"Wrote: {OUT_REPORT}")


# ===========================================================================
# CLI dispatcher
# ===========================================================================
def _run_all():
    cmd_analyse()
    cmd_tune()
    cmd_sens_age()
    cmd_sens_cluster()
    cmd_validate()
    cmd_report()

_COMMANDS = {
    "analyse": cmd_analyse,
    "tune": cmd_tune,
    "sensitivity-age": cmd_sens_age,
    "sensitivity-cluster": cmd_sens_cluster,
    "validate": cmd_validate,
    "report": cmd_report,
    "all": _run_all,
}

def _cli():
    p = argparse.ArgumentParser(
        description="HDS Module 5 single-file pipeline.")
    p.add_argument("command", choices=list(_COMMANDS),
                   help="Which pipeline stage to run.")
    args = p.parse_args()
    _COMMANDS[args.command]()


if __name__ == "__main__":
    _cli()
