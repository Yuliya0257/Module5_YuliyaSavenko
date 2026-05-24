"""
build_report.py
===============
Generate the final Word document for the assignment.
Run after `code/run_analysis.py` so figures are present.
"""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parent
FIG = ROOT / "figures"
OUT = ROOT / "HDS_ML_Savenko_2605.docx"


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------
def style_normal(doc):
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(11)


def add_heading(doc, text, level):
    sizes = {0: 18, 1: 14, 2: 12}
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 11))
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def add_para(doc, text, italic=False, bold=False, size=None,
             alignment=None, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_figure(doc, path, caption, width_cm=14):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cap.add_run(caption)
    cr.italic = True
    cr.font.size = Pt(10)


def add_hyperlink(paragraph, url, text, size=10):
    """Append a clickable external hyperlink run to an existing paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/"
        "relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")        # standard Word hyperlink blue
    rPr.append(color)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))  # half-points
    rPr.append(sz)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


# Blue table styling (ColorBrewer RdBu tints, matching the figures)
TABLE_BORDER = "4393C3"      # medium blue table grid lines
TABLE_HEADER = "D1E5F0"      # light blue header-row shading


def _set_table_borders(table, color=TABLE_BORDER, sz=6):
    """Set every table grid line (outer + inner) to a single colour."""
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))      # eighths of a point
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    # Insert in the schema-correct position within tblPr
    tblPr.insert_element_before(
        borders, "w:shd", "w:tblLayout", "w:tblCellMar",
        "w:tblLook", "w:tblCaption", "w:tblDescription")


def _shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def add_table(doc, headers, rows, widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    _set_table_borders(table)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        _shade_cell(cell, TABLE_HEADER)
        for r in cell.paragraphs[0].runs:
            r.bold = True
    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            table.rows[r_idx].cells[c_idx].text = str(val)
    return table


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------
doc = Document()
style_normal(doc)

# ---- Title page -----------------------------------------------------------
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run(
    "Predicting Cardiovascular Risk Profiles from Depressive "
    "Symptoms in Younger Adults: A Machine Learning Analysis "
    "of NHANES Data"
)
r.bold = True
r.font.size = Pt(16)

add_para(doc, "", space_after=0)
add_para(doc, "Yuliya Savenko",
         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=12)
add_para(doc, "4 June 2026",
         alignment=WD_ALIGN_PARAGRAPH.CENTER, size=11)
add_para(doc, "GitHub: https://github.com/Yuliya0257/Module5_YuliyaSavenko",
         alignment=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=10)

doc.add_paragraph()

# ---- Abstract -------------------------------------------------------------
add_heading(doc, "Abstract", 1)
add_para(doc,
    "Background: Depression in young adulthood is rising and is an "
    "established cardiovascular disease (CVD) risk factor, yet the "
    "evidence base draws predominantly from older populations. "
    "Whether depressive symptomatology improves CVD risk "
    "discrimination beyond demographics in adults under 40 is "
    "unclear. We hypothesised that the Patient Health "
    "Questionnaire-9 (PHQ-9) would add reproducible discriminative "
    "value over demographics alone.")
add_para(doc,
    "Methods: We analysed the NHANES pre-pandemic cycle (2017–March "
    "2020), restricting to adults aged 18–39 with a usable PHQ-9 "
    "(n = 2,728). Data-driven cardiovascular risk profiles were "
    "derived by K-means clustering of 13 self-reported and "
    "behavioural risk indicators. Three classifiers (logistic "
    "regression, random forest and gradient boosting) predicted "
    "profile membership from PHQ-9 plus demographics, each versus "
    "a demographics-only model. Discrimination used 5-fold "
    "stratified cross-validation over three seeds, compared by "
    "paired t-test; sensitivity analyses examined age, clustering "
    "method and fairness.")
add_para(doc,
    "Results: Clustering identified two profiles (silhouette 0.37); "
    "the elevated-risk profile (22.0%) was marked by "
    "physician-recorded lifestyle advice and self-reported "
    "hypertension and dyslipidaemia. Its prevalence was 35.3% with "
    "PHQ-9 ≥ 10 versus 20.6% below 10. Adding PHQ-9 raised "
    "cross-validated ROC-AUC from 0.602 to 0.637 (logistic "
    "regression), 0.592 to 0.638 (random forest) and 0.599 to "
    "0.642 (gradient boosting); every increment was significant "
    "(all p < 10⁻⁵), and PHQ-9 ranked highest by "
    "permutation importance.")
add_para(doc,
    "Conclusions: Depressive symptomatology provides a small but "
    "statistically robust, algorithm-independent improvement in "
    "predicting cardiovascular risk profiles in young adults. The "
    "modest discrimination (AUC ≈ 0.64) indicates the PHQ-9 should "
    "complement, not replace, structured cardiovascular assessment. "
    "Code and dataset-access instructions are openly available.")

doc.add_page_break()

# ---- Introduction ---------------------------------------------------------
add_heading(doc, "1. Introduction", 1)
add_para(doc,
    "Depression is among the leading contributors to the global "
    "disease burden, and its prevalence in young people has risen "
    "sharply over the past decade (Thapar et al., 2022). Large "
    "epidemiological studies have established depression as an "
    "independent cardiovascular disease (CVD) risk factor: the "
    "American Heart Association’s Life’s Essential 8 was recently "
    "extended to Life’s Crucial 9 with a psychological-health "
    "dimension, motivated by NHANES evidence linking depression to "
    "cardiovascular mortality (Ge et al., 2024; Xu et al., 2024). "
    "Subsequent analyses show depression severity tracks "
    "cardiovascular–kidney–metabolic (CKM) syndrome progression (He "
    "et al., 2025; Wang et al., 2025), mortality in diabetes (Zhang Z "
    "et al., 2025), and is mediated in part by inflammation (Meng et "
    "al., 2025).")

add_para(doc,
    "Most of this evidence draws on older adults. Younger adults "
    "accumulate sub-clinical risk for years before events occur, so "
    "risk stratification must rely on modifiable factors and family "
    "history. Machine learning (ML) is well suited here, capturing "
    "non-linear interactions and producing data-driven phenotypes "
    "classical risk equations may miss (Dipnall et al., 2016). "
    "NHANES has become a workhorse for such analyses (Zhang YZ et al., "
    "2025; Vu et al., 2025; Wu et al., 2025; Liu et al., 2025), and "
    "lifestyle and metabolic correlates of depression are well "
    "documented in it (Liu et al., 2024; Zhou et al., 2024; Guo et "
    "al., 2025; Tian et al., 2025).")

add_para(doc,
    "What remains under-explored is whether the routinely-available "
    "PHQ-9 can predict a data-driven CV risk phenotype specifically "
    "in young adults. We therefore asked: can depression in younger "
    "adults be used to predict cardiovascular risk profiles using "
    "machine learning? In NHANES adults aged 18–39 we derived "
    "unsupervised CV risk profiles, trained three classifiers to "
    "predict profile membership from PHQ-9 plus demographics, and "
    "compared against a demographics-only baseline, hypothesising "
    "that depression would offer incremental discriminative value.")

# ---- Methods --------------------------------------------------------------
add_heading(doc, "2. Methods", 1)

add_heading(doc, "2.1 Data source and rationale", 2)
add_para(doc,
    "NHANES is a continuous multi-stage stratified probability "
    "sample of the US non-institutionalised civilian population. "
    "Three features make it well-suited to ML for early-adulthood CV "
    "research: (i) it couples the validated PHQ-9 with self-reported "
    "diagnoses, medication and physiological measurements in the "
    "same participants; (ii) the per-cycle sample (≈ 15,000 overall; "
    "≈ 3,000 aged 18–39) supports repeated stratified CV and stable "
    "feature importance; (iii) every variable is documented in a "
    "public codebook with stable cross-cycle coding, supporting "
    "reproducibility, a benchmark met by the NHANES-ML literature "
    "(Vu et al., 2025; Wu et al., 2025; Zhang YZ et al., 2025).")

add_para(doc,
    "NHANES uses stratified multi-stage cluster sampling with "
    "oversampling of minority and low-income strata. We treat it as "
    "a learning sample and report unweighted ML metrics, since "
    "accuracy, F1 and AUC describe the model on data from this "
    "design rather than a simple random sample; the minority "
    "oversampling aids fairness audits but means sample prevalence "
    "≠ national prevalence. We used the 2017–March 2020 pre-"
    "pandemic combined cycle, the most recent NHANES release "
    "containing PHQ-9 alongside the cardiovascular and medical-"
    "conditions modules.")

add_heading(doc, "2.2 Sample", 2)
add_para(doc,
    "Three considerations motivated the 18–39 window: CV risk "
    "accumulates at this stage before events, when modifiable "
    "factors remain amenable to intervention; depressive-symptom "
    "prevalence peaks in early adulthood (Thapar et al., 2022); and "
    "restricting to younger adults reduces confounding by "
    "comorbidities and polypharmacy.")

add_para(doc,
    "From 3,260 participants aged 18–39 we excluded pregnant "
    "respondents and those without a usable PHQ-9 (≤ 2 items "
    "missing); n = 2,728.")

add_heading(doc, "2.3 Variables", 2)
add_para(doc,
    "Depression. The PHQ-9 is a nine-item self-report scale (total "
    "0–27); PHQ-9 ≥ 10 is the conventional threshold for clinically "
    "significant symptoms. We modelled PHQ-9 as continuous in the "
    "supervised models and binary in cross-tabulations.")

add_para(doc,
    "Cardiovascular risk indicators (13): self-reported hypertension "
    "(BPQ020), antihypertensive use (BPQ050A), high cholesterol "
    "(BPQ080), lipid-lowering use (BPQ100D), doctor-reported "
    "overweight (MCQ080), four physician lifestyle directives "
    "(MCQ366A–D), family history of premature MI (MCQ300A) and "
    "diabetes (MCQ300C), insufficient sleep (< 7 h) and household "
    "tobacco exposure. Because clinical events are rare before age "
    "40, physician-recognised risk factors provide a more "
    "informative substrate.")

add_para(doc,
    "Demographics and socioeconomic status. Age, sex, race/"
    "ethnicity, education, marital status, and family poverty-to-"
    "income ratio.")

add_heading(doc, "2.4 Pre-processing", 2)
add_para(doc,
    "NHANES refused/don’t-know codes were set to missing; binary "
    "indicators were mode-imputed, continuous indicators median-"
    "imputed. Categorical variables were one-hot encoded; continuous "
    "variables z-score scaled, with scaling parameters estimated on "
    "the training fold only to avoid leakage.")

add_heading(doc, "2.5 Cardiovascular risk-profile derivation", 2)
add_para(doc,
    "Rationale for clustering rather than a published score: (i) "
    "calibrated scores (Framingham, AHA PREVENT, QRISK3) were "
    "derived on adults ≥ 30–40 and collapse near “low risk” here; "
    "(ii) they require measured biomarkers unavailable in our "
    "subset; (iii) a fixed-weight composite ignores risk-factor "
    "co-occurrence patterns recognised in cardiovascular-kidney-"
    "metabolic phenotyping (He et al., 2025; Wang et al., 2025).")

add_para(doc,
    "K-means (k-means++, 20 restarts) on the standardised indicators "
    "recovers the constellation in which physician-flagged lifestyle "
    "issues, diagnoses and family history co-occur, a learned "
    "“clinician-recognised at-risk” phenotype (Figure 1B) treated as "
    "an analytic target, not a clinical diagnosis (Section 4.4). K "
    "was selected by mean silhouette over K ∈ {2,3,4}.")

add_heading(doc, "2.6 Supervised models", 2)
add_para(doc,
    "Three classifiers with deliberately different inductive biases "
    "were used: multinomial logistic regression (L2-regularised, "
    "linear, with directly interpretable log-odds coefficients); a "
    "random forest (bagged CART trees); and gradient boosting. "
    "Models were implemented in Python; full code is in the GitHub "
    "repository. For each algorithm we fit a demographics-only "
    "baseline alongside the full model: a signal surviving linear, "
    "bagged and boosted models alike is unlikely to be a single-"
    "model artefact.")

add_heading(doc, "2.7 Hyperparameter tuning", 2)
add_para(doc,
    "Hyperparameters were selected by grid search with 3-fold "
    "stratified CV on an 80% tuning subset (n = 2,182). Selection "
    "criterion: mean macro F1. The LR grid spanned 16 configs, RF "
    "27 and GB 12 (Table 2); the full 55-config grid is in "
    "results/hyperparameter_search.csv.")

add_table(doc,
    headers=["Algorithm", "Hyperparameter", "Values tested", "Selected"],
    rows=[
        ["Logistic regression", "C (inverse L2)", "0.01, 0.1, 1, 10", "1"],
        ["Logistic regression", "Learning rate", "0.1, 0.3, 0.5, 1.0", "1.0"],
        ["Random forest", "n_estimators", "40, 60, 100", "60"],
        ["Random forest", "max_depth", "4, 6, 8", "8"],
        ["Random forest", "min_samples_leaf", "20, 30, 50", "30"],
        ["Gradient boosting", "n_estimators", "40, 60, 100", "60"],
        ["Gradient boosting", "Learning rate", "0.05, 0.1", "0.05"],
        ["Gradient boosting", "max_depth", "2, 3", "3"],
    ],
)
add_para(doc, "Table 2. Hyperparameter search space and selected "
              "configurations (√p features per split fixed for the "
              "tree ensembles).",
         italic=True, size=10)

add_heading(doc, "2.8 Evaluation and interpretability", 2)
add_para(doc,
    "Performance was estimated by stratified 5-fold CV over three "
    "seeds (15 folds), with minority-class oversampling on training "
    "folds only. We report accuracy, balanced accuracy, macro F1 and "
    "macro-averaged one-versus-rest ROC-AUC; the full-vs-baseline "
    "macro F1 difference was tested with a paired t-test. On a "
    "held-out 25% split we computed RF permutation importance, LR "
    "coefficients, subgroup performance (sex, race/ethnicity, income "
    "tertile), and calibration (Brier score, predicted vs observed "
    "prevalence).")

add_heading(doc, "2.9 Software and ethics", 2)
add_para(doc,
    "Python 3.10, NumPy 2.2, pandas 2.3, matplotlib 3.10, seaborn "
    "0.13; reproducible from code/run_analysis.py. NHANES is "
    "public-domain de-identified data; IRB approval was not "
    "required.")

# ---- Results --------------------------------------------------------------
add_heading(doc, "3. Results", 1)

add_heading(doc, "3.1 Cohort", 2)
add_para(doc,
    "The analytic cohort comprised 2,728 adults aged 18–39 (mean age "
    "28.0; 51.0% female; 9.2% with PHQ-9 ≥ 10). Table 1 summarises "
    "characteristics by derived risk profile; every cardiovascular "
    "indicator and the PHQ-9 is higher in the elevated-risk "
    "profile.")

add_table(doc,
    headers=["Characteristic", "Overall", "Low-risk profile",
             "Elevated-risk profile"],
    rows=[
        ["n", "2,728", "2,128", "600"],
        ["Age, mean (SD)", "28.0 (6.6)", "27.6 (6.5)", "29.5 (6.5)"],
        ["Female, %", "51.0", "48.6", "59.7"],
        ["Non-Hispanic White, %", "31.3", "31.7", "29.7"],
        ["Non-Hispanic Black, %", "25.0", "25.0", "25.0"],
        ["Mexican American, %", "14.8", "13.8", "18.5"],
        ["Non-Hispanic Asian, %", "12.1", "12.8", "9.7"],
        ["PHQ-9 total, mean (SD)", "3.4 (4.2)", "3.0 (3.9)", "4.5 (4.8)"],
        ["PHQ-9 ≥ 10, %", "9.2", "7.7", "14.8"],
        ["Self-reported hypertension, %", "12.1", "7.5", "28.3"],
        ["Self-reported high cholesterol, %", "10.6", "5.5", "28.7"],
        ["Doctor-reported overweight, %", "33.7", "19.1", "85.3"],
        ["Family history, premature MI, %", "8.6", "7.2", "13.3"],
        ["Family history, diabetes, %", "34.6", "30.7", "48.3"],
        ["Insufficient sleep (< 7 h), %", "23.7", "23.3", "25.3"],
        ["Household smoker, %", "35.6", "36.3", "33.0"],
    ],
)
add_para(doc, "Table 1. Cohort characteristics, overall and by "
              "derived cardiovascular risk profile (n = 2,728).",
         italic=True, size=10)

add_heading(doc, "3.2 Cardiovascular risk profiles", 2)
add_para(doc,
    "K-means selected K = 2 (silhouette 0.373 versus 0.274 and 0.285 "
    "for K = 3 and K = 4; Figure 1A). The two profiles separate "
    "sharply on physician-given lifestyle advice (Figure 1B): in "
    "profile 1 (n = 600, 22.0%), standardised scores for being told "
    "to lose weight, exercise, reduce salt and reduce fat ranged "
    "+1.12 to +1.58 SD versus −0.44 to −0.31 SD in profile 0, with "
    "elevated self-reported hypertension (+0.50 SD) and "
    "hypercholesterolaemia (+0.59 SD). We label profile 0 “low” and "
    "profile 1 “elevated” CV risk.")

add_figure(doc, FIG / "fig_silhouette.png",
           "Figure 1A. Silhouette and elbow plots for K-means cluster "
           "selection over K ∈ {2,3,4}.", width_cm=14)
add_figure(doc, FIG / "fig_cluster_profile.png",
           "Figure 1B. Standardised K-means centroids on cardiovascular "
           "risk indicators. The elevated-risk profile is dominated by "
           "physician-given lifestyle directives (MCQ366A–D) and "
           "self-reported hypertension and dyslipidaemia.", width_cm=15)

add_heading(doc, "3.3 Crude depression–risk-profile association", 2)
add_para(doc,
    "The elevated-risk profile contained 35.3% of participants with "
    "PHQ-9 ≥ 10 versus 20.6% of those with PHQ-9 < 10: clinically "
    "significant depression doubled the odds of the elevated-risk "
    "profile (odds ratio 2.10, 95% CI 1.59–2.77; χ² = 28.7, "
    "p < 0.001). Median PHQ-9 was 2 in the low-risk profile versus "
    "3 in the high-risk profile (Figure 2).")

add_figure(doc, FIG / "fig_phq_by_profile.png",
           "Figure 2. Distribution of PHQ-9 total scores by cardiovascular "
           "risk profile. Box plot summarises the inter-quartile range; "
           "scattered points show a sub-sample of individual "
           "participants.", width_cm=16)

add_heading(doc, "3.4 Model performance", 2)
add_para(doc,
    "Cross-validated discrimination is in Table 3 and Figure 3. "
    "Adding PHQ-9 raised ROC-AUC by ≈ 0.04 in all three algorithms "
    "(LR: 0.602 → 0.637; RF: 0.592 → 0.638; GB: 0.599 → 0.642), the "
    "boosted ensemble marginally strongest, with good stability "
    "across folds (SDs 0.02–0.04). Paired t-tests confirmed "
    "significant macro-F1 improvement under every algorithm (LR "
    "t = 4.49; RF t = 7.47; GB t = 5.48; all p < 10⁻⁵).")

add_table(doc,
    headers=["Model", "Accuracy", "Balanced acc", "Macro F1", "ROC-AUC"],
    rows=[
        ["Logistic regression (demographics only)",
         "0.575 ± 0.020", "0.573 ± 0.022", "0.525 ± 0.018", "0.602 ± 0.026"],
        ["Logistic regression (+ PHQ-9)",
         "0.605 ± 0.020", "0.591 ± 0.029", "0.547 ± 0.022", "0.637 ± 0.031"],
        ["Random forest (demographics only)",
         "0.608 ± 0.022", "0.565 ± 0.021", "0.536 ± 0.019", "0.592 ± 0.031"],
        ["Random forest (+ PHQ-9)",
         "0.634 ± 0.022", "0.604 ± 0.027", "0.568 ± 0.023", "0.638 ± 0.034"],
        ["Gradient boosting (demographics only)",
         "0.592 ± 0.026", "0.571 ± 0.028", "0.532 ± 0.024", "0.599 ± 0.034"],
        ["Gradient boosting (+ PHQ-9)",
         "0.620 ± 0.025", "0.606 ± 0.035", "0.562 ± 0.028", "0.642 ± 0.036"],
    ],
)
add_para(doc, "Table 3. Cross-validated performance (mean ± SD, "
              "5-fold × 3 seeds). Adding PHQ-9 improves every model.",
         italic=True, size=10)

add_figure(doc, FIG / "fig_cv_metrics.png",
           "Figure 3. Cross-validated performance of the six models "
           "(three algorithms × full / demographics-only). Adding "
           "PHQ-9 consistently improves on the demographics-only "
           "baseline.", width_cm=15)

add_heading(doc, "3.5 Confusion and interpretability", 2)
add_para(doc,
    "On the held-out 25% split the RF confusion matrix (Figure 4A) "
    "showed moderate sensitivity for the elevated-risk class "
    "(≈ 0.53), with false positives expected from training-time "
    "oversampling (Section 4.3).")

add_figure(doc, FIG / "fig_rf_confusion.png",
           "Figure 4A. Random forest confusion matrix on the held-out 25% "
           "test set. Cell shading shows row-normalised proportion; "
           "numerals are raw counts.", width_cm=10)

add_para(doc,
    "Permutation importance ranked PHQ-9 as the single most "
    "informative feature (Figure 4B; macro-F1 drop 0.032 ± 0.013), "
    "well ahead of any other variable; LR coefficients corroborated "
    "this (PHQ-9 +0.19 standardised, comparable to age +0.21). The "
    "marginal RF effect of PHQ-9 (Figure 5) was monotonic, rising "
    "from ≈ 0.48 at PHQ-9 = 0 to ≈ 0.65 at PHQ-9 = 20.")

add_figure(doc, FIG / "fig_rf_permimp.png",
           "Figure 4B. Random forest permutation importance "
           "(8 permutations, macro F1 drop). PHQ-9 total is the "
           "single most informative feature.", width_cm=14)

add_figure(doc, FIG / "fig_rf_phq_marginal.png",
           "Figure 5. Marginal random-forest predicted probability of "
           "each risk profile as a function of PHQ-9 total, with "
           "demographics held at sample means.", width_cm=13)

add_heading(doc, "3.6 Sensitivity analyses", 2)
add_para(doc,
    "Age window. Re-running on broader age windows (Table 4) showed "
    "the incremental PHQ-9 ROC-AUC was largest in 18–39 (+0.036 LR / "
    "+0.042 RF) and roughly halved in 18–64 and 18+. Baseline AUC "
    "rose with age range as accumulated risk factors carry more "
    "weight in older adults; all ablations stayed significant "
    "(p < 0.005). Depression's marginal value is therefore greatest "
    "in the under-40 window.")

add_para(doc,
    "Clustering method. Re-deriving the profiles with K-modes (the "
    "matching-dissimilarity method appropriate for binary data) "
    "agreed closely with the K-means profiles (98.3% agreement, "
    "adjusted Rand index 0.92) and preserved the depression gradient "
    "(31.7% elevated-risk in PHQ-9 ≥ 10 vs 19.1% otherwise). The "
    "phenotype is not an artefact of K-means on binary data.")

add_table(doc,
    headers=["Cohort", "n", "Elev %", "LR demo AUC", "LR + PHQ AUC", "ΔAUC", "RF demo AUC", "RF + PHQ AUC", "ΔAUC"],
    rows=[
        ["18–39 (primary)", "2,728", "22.0", "0.602", "0.637", "+0.036", "0.595", "0.637", "+0.042"],
        ["18–64 (working-age)", "6,254", "34.5", "0.668", "0.685", "+0.017", "0.675", "0.691", "+0.016"],
        ["18+ (all adults)",  "8,223", "36.7", "0.654", "0.677", "+0.024", "0.671", "0.692", "+0.021"],
    ],
)
add_para(doc, "Table 4. Age-window sensitivity analysis (5-fold CV × "
              "2 seeds). PHQ-9’s incremental discriminative value is "
              "largest in the under-40 cohort.",
         italic=True, size=10)

add_heading(doc, "3.7 Subgroup performance (fairness check)", 2)
add_para(doc,
    "Held-out RF subgroup performance is in Table 5. Sex parity was "
    "good (AUC 0.608 vs 0.606); race/ethnicity AUC ranged 0.586 to "
    "0.631, balanced accuracy lowest in Mexican American (0.511); "
    "income tertiles showed a modest gradient (AUC 0.598 → 0.622).")

add_table(doc,
    headers=["Subgroup", "n", "% elevated", "Bal acc", "Macro F1", "ROC-AUC"],
    rows=[
        ["Overall", "682", "22.0", "0.580", "0.543", "0.609"],
        ["Sex: Female", "354", "23.7", "0.568", "0.500", "0.608"],
        ["Sex: Male", "328", "20.1", "0.577", "0.568", "0.606"],
        ["Race: Mexican American", "102", "26.5", "0.511", "0.470", "0.587"],
        ["Race: Non-Hispanic Asian", "90", "11.1", "0.544", "0.460", "0.586"],
        ["Race: Non-Hispanic Black", "162", "25.9", "0.611", "0.589", "0.631"],
        ["Race: Non-Hispanic White", "209", "23.4", "0.586", "0.561", "0.613"],
        ["Race: Other Hispanic", "80", "22.5", "0.637", "0.623", "0.617"],
        ["Income (PIR) low <1.47", "227", "20.7", "0.576", "0.550", "0.598"],
        ["Income (PIR) mid 1.47–2.46", "226", "22.1", "0.570", "0.544", "0.610"],
        ["Income (PIR) high ≥2.46", "229", "23.1", "0.588", "0.526", "0.622"],
    ],
)
add_para(doc, "Table 5. Held-out RF subgroup performance "
              "(n = 682 test split). The Other/Multi-race subgroup "
              "(n = 39) is omitted because its small size renders "
              "metrics unstable.",
         italic=True, size=10)

# ---- Discussion -----------------------------------------------------------
add_heading(doc, "4. Discussion", 1)

add_heading(doc, "4.1 Summary of key findings", 2)
add_para(doc,
    "The results broadly support the hypothesis that depressive "
    "symptomatology carries cardiovascular-risk information in young "
    "adults beyond what demographics provide. Even in an under-40 "
    "sample where overt CVD is rare, the PHQ-9 added a small but "
    "reproducible improvement to discrimination of a data-driven CV "
    "risk phenotype. The most informative finding is the consistency "
    "of the signal: three algorithms with different inductive biases "
    "converged on ROC-AUC ≈ 0.64, every ablation was significant, "
    "and PHQ-9 ranked as the most important input by permutation "
    "importance. Crude elevated-risk prevalence was 71% higher among "
    "PHQ-9 ≥ 10 participants (35.3% vs 20.6%).")

add_para(doc,
    "These findings align with the accumulating NHANES literature "
    "linking depression to cardiovascular health (Ge et al., 2024; "
    "Xu et al., 2024; Wu et al., 2025; He et al., 2025; Wang et al., "
    "2025; Zhang Z et al., 2025; Meng et al., 2025). Our contribution "
    "is to show the signal is recoverable from PHQ-9 alone against a "
    "learned multivariate phenotype, with incremental value greatest "
    "in the under-40 window.")

add_heading(doc, "4.2 Model comparison and trade-offs", 2)
add_para(doc,
    "All three algorithms converged on near-identical discrimination "
    "(ROC-AUC 0.637–0.642), itself informative: the depression–CV-"
    "risk signal is largely additive and linearly accessible, with "
    "little extra recoverable from non-linear interactions. This "
    "shapes the model-choice trade-off. Logistic regression is the "
    "most transparent (coefficients read as log-odds, fast, simple "
    "to audit and recalibrate); the tree ensembles capture "
    "interactions but at the cost of interpretability (needing post-"
    "hoc tools such as SHAP) and greater computational expense. For "
    "a primary-care screening application, where transparency "
    "matters and the boosted model's 0.005 AUC edge is not "
    "clinically meaningful, logistic regression is preferable; for "
    "a research setting prioritising predictive ceiling, gradient "
    "boosting would be chosen. The pipeline retains all three.")

add_heading(doc, "4.3 Calibration and clinical utility", 2)
add_para(doc,
    "Discrimination and calibration are distinct, and the headline "
    "metric concerns only the former. An AUC of 0.64 falls below the "
    "0.70 conventionally taken as the minimum for acceptable "
    "discrimination: given one elevated-risk and one low-risk "
    "participant, the model ranks them correctly 64% of the time. "
    "Calibration is weaker still. Because training folds were "
    "oversampled to 50:50, the random forest's predicted "
    "probabilities are inflated (mean 0.46 against a true prevalence "
    "of 0.22) and its Brier score (0.229) is worse than simply "
    "predicting the base rate (0.172); a non-oversampled model is "
    "well calibrated (mean 0.22, Brier 0.166) but loses minority-"
    "class sensitivity. The probabilities in Figure 5 should "
    "therefore be read as relative ordering, not true risks. "
    "Clinically, this is not a deployable individual risk "
    "calculator. Its defensible uses are narrower: a low-stakes "
    "triage signal, where even weak discrimination can carry net "
    "benefit because flagging a young adult for fuller "
    "cardiovascular assessment is cheap and low-harm, and "
    "population-level evidence. It should not inform treatment "
    "decisions.")

add_heading(doc, "4.4 Strengths and limitations", 2)
add_para(doc,
    "The analysis has several strengths: a strict ablation against a "
    "demographics-only baseline, documented grid-search tuning, and "
    "age-window, clustering-method and subgroup-fairness sensitivity "
    "analyses. In keeping with open-science practice the whole "
    "pipeline is fully reproducible, with all code, result files and "
    "dataset-access instructions openly available for independent "
    "re-running and audit. Six limitations qualify the "
    "interpretation. (i) The design is cross-sectional, so reverse "
    "causation remains plausible. (ii) CV indicators are self-"
    "reported, biasing ascertainment toward those with healthcare "
    "access. (iii) Measured biomarkers were unavailable. (iv) "
    "Discrimination (AUC ≈ 0.64) is modest. (v) The K-means "
    "phenotype is an analytic target, not a prospectively validated "
    "risk class; external validation against incident events is "
    "essential before clinical use. (vi) Although NHANES is a "
    "probability sample, the complete-case analytic subset may "
    "under-represent the most survey- and healthcare-disengaged "
    "young adults, modestly limiting external validity. A minor "
    "redundancy is also noted: sleep enters the clustering through "
    "two correlated indicators, mildly over-weighting it, though its "
    "near-zero cluster separation (Figure 1B) makes the effect "
    "negligible.")

add_heading(doc, "4.5 Ethics, fairness and interpretability", 2)
add_para(doc,
    "The subgroup audit (Table 5) shows good sex parity (AUC 0.61 "
    "vs 0.61) but uneven race/ethnicity discrimination, weakest in "
    "Mexican American (AUC 0.59, balanced accuracy 0.51) and Non-"
    "Hispanic Asian (AUC 0.59) participants, with a modest income "
    "gradient (AUC 0.60 → 0.62). These disparities reflect "
    "underlying differences in CV-risk prevalence and clinician "
    "flagging rather than biological causation, but indicate the "
    "model would under-serve some subgroups in deployment. The "
    "reliance on self-reported diagnoses compounds this: groups "
    "with poorer healthcare access are under-diagnosed and so "
    "receive systematically lower risk estimates, risking under-"
    "recognition of those most in need. Clinical-translation "
    "pipelines should remove race or replace it with social-"
    "determinant-of-health proxies (Vu et al., 2025; Wu et al., "
    "2025). Using depression as a CV risk marker also risks stigma "
    "for high-PHQ-9 patients; deployment would require governance "
    "and fairness audits.")

add_heading(doc, "4.6 Implications and future development", 2)
add_para(doc,
    "The findings support integrating mental-health screening into "
    "early-adulthood CV cohorts. Future work should pool additional "
    "NHANES cycles to enlarge the under-40 sample, validate the "
    "phenotype against incident cardiovascular events, and re-derive "
    "the risk profiles on measured biomarkers.")

# ---- Conclusion -----------------------------------------------------------
add_heading(doc, "5. Conclusion", 1)
add_para(doc,
    "Overall, the findings tell a coherent story. In NHANES adults "
    "aged 18–39, depressive symptomatology measured by the PHQ-9 "
    "contributes a small but reproducible improvement over "
    "demographics in predicting a data-driven cardiovascular risk "
    "phenotype, consistent across three algorithms and largest in "
    "the under-40 window. The effect is modest and should not be "
    "overinterpreted: PHQ-9 is not a stand-alone risk tool, but it "
    "carries enough signal to support depression screening in "
    "early-adulthood CV risk assessment, reinforcing the view that "
    "psychological health is part of cardiovascular health.")

# ---- References -----------------------------------------------------------
add_heading(doc, "References", 1)
# Harvard-style references: alphabetical, numbered, with the journal
# name italicised. Each entry is (authors + year + title, journal,
# volume/pages, PubMed URL). Eleven links use the article's verified
# PMID; five recent articles, whose PMIDs could not be confirmed, use
# a PubMed DOI lookup that resolves to the same record.
refs = [
    ("Dipnall, J.F., Pasco, J.A., Berk, M. et al. (2016) 'Into the "
     "Bowels of Depression: Unravelling Medical Symptoms Associated "
     "with Depression by Applying Machine-Learning Techniques to a "
     "Community Based Population Sample', ",
     "PLoS One",
     ", 11(12), e0167055.",
     "https://pubmed.ncbi.nlm.nih.gov/27935995/"),

    ("Ge, J., Peng, W. and Lu, J. (2024) 'Predictive Value of Life’s "
     "Crucial 9 for Cardiovascular and All-Cause Mortality: A "
     "Prospective Cohort Study From the NHANES 2007 to 2018', ",
     "Journal of the American Heart Association",
     ", 13(19), e036669.",
     "https://pubmed.ncbi.nlm.nih.gov/39377201/"),

    ("Guo, T., Zhou, Y., Yang, G., Sheng, L. and Chai, X. (2025) "
     "'Association between cardiometabolic index and hypertension "
     "among US adults from NHANES 1999–2020', ",
     "Scientific Reports",
     ", 15, 87029.",
     "https://pubmed.ncbi.nlm.nih.gov/?term=10.1038%2Fs41598-025-87029-0"),

    ("He, Y., Lan, L., Liu, Y. et al. (2025) 'Depression severity as "
     "a predictor of cardiovascular-kidney-metabolic syndrome "
     "progression and mortality: Results from two nationally "
     "representative cohort studies', ",
     "Journal of Affective Disorders",
     ", 388, 119606.",
     "https://pubmed.ncbi.nlm.nih.gov/?term=10.1016%2Fj.jad.2025.119606"),

    ("Liu, H., Dong, H., Zhou, Y. et al. (2024) 'The association "
     "between Metabolic Score for Visceral Fat and depression in "
     "overweight or obese individuals', ",
     "Frontiers in Endocrinology",
     ", 15, 1482003.",
     "https://pubmed.ncbi.nlm.nih.gov/?term=10.3389%2Ffendo.2024.1482003"),

    ("Liu, X., Luo, Z., Jing, F. et al. (2025) 'Estimating "
     "cardiovascular mortality in patients with hypertension using "
     "machine learning: the role of depression classification based "
     "on lifestyle and physical activity', ",
     "Journal of Psychosomatic Research",
     ", 189, 112030.",
     "https://pubmed.ncbi.nlm.nih.gov/39752763/"),

    ("Meng, X., Han, L., Fu, J., Hu, C. and Lu, Y. (2025) "
     "'Associations between metabolic syndrome and depression, and "
     "the mediating role of inflammation: Based on the NHANES "
     "database', ",
     "Journal of Affective Disorders",
     ", 375, pp. 214–221.",
     "https://pubmed.ncbi.nlm.nih.gov/39862983/"),

    ("Thapar, A., Eyre, O., Patel, V. and Brent, D. (2022) "
     "'Depression in young people', ",
     "The Lancet",
     ", 400(10352), pp. 617–631.",
     "https://pubmed.ncbi.nlm.nih.gov/35940184/"),

    ("Tian, Q., Guo, J., Ding, J. and Zhu, Y. (2025) 'The "
     "relationship between unhealthy lifestyle behaviors and "
     "depression: Evidence from NHANES', ",
     "Journal of Affective Disorders",
     ", 384, pp. 214–221.",
     "https://pubmed.ncbi.nlm.nih.gov/40258422/"),

    ("Vu, T., Dawadi, R., Yamamoto, M. et al. (2025) 'Prediction of "
     "depressive disorder using machine learning approaches: "
     "findings from the NHANES', ",
     "BMC Medical Informatics and Decision Making",
     ", 25, 83.",
     "https://pubmed.ncbi.nlm.nih.gov/39962516/"),

    ("Wang, G., Wu, Y., Chen, A. et al. (2025) 'Association between "
     "depressive symptoms and mortality in patients with "
     "Cardiovascular-Kidney-Metabolic syndrome: the mediating role "
     "of inflammatory biomarkers', ",
     "Journal of Affective Disorders",
     ", 386, 119429.",
     "https://pubmed.ncbi.nlm.nih.gov/?term=10.1016%2Fj.jad.2025.119429"),

    ("Wu, Z., Xu, P., Zhai, Y. et al. (2025) 'The Association of "
     "Elevated Depression Levels and Life’s Essential 8 on "
     "Cardiovascular Health With Predicted Machine Learning Models "
     "and Interpretations: Evidence From NHANES 2007–2018', ",
     "Depression and Anxiety",
     ", 2025, 8865176.",
     "https://pubmed.ncbi.nlm.nih.gov/40255861/"),

    ("Xu, Y., Ning, W., Zhang, Y. et al. (2024) 'Associations "
     "Between Cardiovascular Health (Life’s Essential 8) and Mental "
     "Disorders', ",
     "Clinical Cardiology",
     ", 47(10), e70019.",
     "https://pubmed.ncbi.nlm.nih.gov/39314085/"),

    ("Zhang, Y.-Z., Wu, H.-Y., Ma, R.-W. et al. (2025) 'Machine "
     "Learning-Based predictive model for adolescent metabolic "
     "syndrome: Utilizing data from NHANES 2007–2016', ",
     "Scientific Reports",
     ", 15, 88156.",
     "https://pubmed.ncbi.nlm.nih.gov/39863763/"),

    ("Zhang, Z., Xu, H., Meng, Y., Yan, Y. and Wang, Y. (2025) "
     "'Association between depressive symptoms and all-cause and "
     "cardiovascular mortality in diabetes patients: A national "
     "cohort study', ",
     "Journal of Affective Disorders",
     ", 391, 119957.",
     "https://pubmed.ncbi.nlm.nih.gov/?term=10.1016%2Fj.jad.2025.119957"),

    ("Zhou, X., Tao, X.-L., Zhang, L. et al. (2024) 'Association "
     "between cardiometabolic index and depression: National Health "
     "and Nutrition Examination Survey (NHANES) 2011–2014', ",
     "Journal of Affective Disorders",
     ", 351, pp. 939–947.",
     "https://pubmed.ncbi.nlm.nih.gov/38341157/"),
]
for i, (part1, journal, part2, url) in enumerate(refs, start=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.first_line_indent = Cm(-0.8)
    for seg, ital in ((f"{i}. ", False), (part1, False),
                      (journal, True), (part2 + "  ", False)):
        r = p.add_run(seg)
        r.font.size = Pt(10)
        r.italic = ital
    add_hyperlink(p, url, "[PubMed]", size=10)

doc.save(str(OUT))
print(f"Wrote: {OUT}")
